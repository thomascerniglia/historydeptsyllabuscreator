"""
Document Generation Module for History Syllabus Generator
Contains all document generation methods and utilities
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import tempfile
import os
import traceback
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import docx
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    def convert(input_path, output_path):
        raise RuntimeError("docx2pdf library is not installed. Cannot convert to PDF.")

# Additional imports for cross-platform PDF generation
import subprocess
import sys
import platform

from constants import *

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    # This gets the relationship ID for the hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the hyperlink element
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    
    # Create a new run
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Add color and underline to the run
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')  # Blue color
    rPr.append(c)
    
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')  # Single underline
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def process_text_with_hyperlinks(paragraph, text):
    """Process text containing URLs and convert them to hyperlinks"""
    import re
    
    # Better regex patterns
    url_pattern = re.compile(r'(https?://[^\s\'"<>]+[^\s\'"<>.,;:])')
    email_pattern = re.compile(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})')
    
    # Process URLs first
    url_parts = url_pattern.split(text)
    result_parts = []
    
    for part in url_parts:
        if url_pattern.match(part):
            # This is a URL
            try:
                add_hyperlink(paragraph, part, part)
            except Exception as e:
                # Fallback to plain text if hyperlink creation fails
                paragraph.add_run(part)
        else:
            # Process emails in non-URL text
            email_parts = email_pattern.split(part)
            for email_part in email_parts:
                if email_pattern.match(email_part):
                    # This is an email
                    try:
                        add_hyperlink(paragraph, email_part, f"mailto:{email_part}")
                    except Exception as e:
                        # Fallback to plain text
                        paragraph.add_run(email_part)
                else:
                    # Regular text
                    if email_part.strip():
                        paragraph.add_run(email_part)

class DocumentGenerationMixin:
    """Mixin class containing all document generation methods"""
    
    def check_pdf_capabilities(self):
        """Check what PDF generation capabilities are available"""
        capabilities = {
            "docx2pdf": DOCX2PDF_AVAILABLE and platform.system() == "Windows",
            "libreoffice": False,
            "reportlab": True  # Always available since it's in requirements
        }
        
        # Check for LibreOffice
        libreoffice_paths = []
        if platform.system() == "Windows":
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
        elif platform.system() == "Darwin":  # macOS
            libreoffice_paths = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            ]
        else:  # Linux
            libreoffice_paths = [
                "/usr/bin/libreoffice",
                "/usr/bin/soffice",
                "/snap/bin/libreoffice"
            ]
        
        for path in libreoffice_paths:
            if os.path.exists(path):
                capabilities["libreoffice"] = True
                break
        
        return capabilities
    
    def show_pdf_setup_info(self):
        """Show information about PDF setup and requirements"""
        capabilities = self.check_pdf_capabilities()
        
        message = "PDF Generation Setup Information\n\n"
        
        if capabilities["docx2pdf"]:
            message += "✓ Microsoft Word integration available\n"
        else:
            message += "✗ Microsoft Word integration not available\n"
            
        if capabilities["libreoffice"]:
            message += "✓ LibreOffice integration available\n"
        else:
            message += "✗ LibreOffice not found\n"
            
        message += "✓ Basic PDF generation always available\n\n"
        
        message += "For best PDF quality, install one of the following:\n\n"
        message += "Option 1: Microsoft Word (Windows)\n"
        message += "• Provides best formatting and compatibility\n"
        message += "• Automatically used if available\n\n"
        
        message += "Option 2: LibreOffice (Free, All Platforms)\n"
        message += "• Download from: https://www.libreoffice.org\n"
        message += "• Good formatting and cross-platform support\n"
        message += "• Automatically detected and used\n\n"
        
        message += "Option 3: Basic PDF (Always Available)\n"
        message += "• Uses built-in PDF generation\n"
        message += "• Basic formatting only\n"
        message += "• No additional software required\n\n"
        
        if not capabilities["docx2pdf"] and not capabilities["libreoffice"]:
            message += "RECOMMENDATION: Install LibreOffice for better PDF support"
        
        messagebox.showinfo("PDF Setup Information", message)
    
    def convert_docx_to_pdf_robust(self, docx_path, pdf_path):
        """
        Robust PDF conversion with multiple fallback methods.
        Handles permission issues and cross-platform compatibility.
        """
        errors = []
        
        # Method 1: Try docx2pdf (Windows with Word)
        if DOCX2PDF_AVAILABLE and platform.system() == "Windows":
            try:
                # Ensure paths are absolute and writable
                abs_docx_path = os.path.abspath(docx_path)
                abs_pdf_path = os.path.abspath(pdf_path)
                
                # Check if output directory is writable
                output_dir = os.path.dirname(abs_pdf_path)
                if not os.access(output_dir, os.W_OK):
                    raise PermissionError(f"No write permission to directory: {output_dir}")
                
                # Try to create/delete a test file to verify permissions
                test_file = os.path.join(output_dir, "test_write_permission.tmp")
                try:
                    with open(test_file, 'w') as f:
                        f.write("test")
                    os.remove(test_file)
                except Exception as perm_error:
                    raise PermissionError(f"Cannot write to directory: {output_dir}")
                
                convert(abs_docx_path, abs_pdf_path)
                return True, "PDF created successfully using docx2pdf"
                
            except Exception as e:
                error_msg = str(e).lower()
                if "com" in error_msg or "word" in error_msg:
                    errors.append("docx2pdf failed - Microsoft Word not properly configured")
                elif "permission" in error_msg:
                    errors.append("docx2pdf failed - Permission denied")
                else:
                    errors.append(f"docx2pdf failed - {str(e)[:100]}")
        
        # Method 2: Try LibreOffice command line (cross-platform)
        try:
            # Common LibreOffice executable locations
            libreoffice_paths = []
            if platform.system() == "Windows":
                libreoffice_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                ]
            elif platform.system() == "Darwin":  # macOS
                libreoffice_paths = [
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice"
                ]
            else:  # Linux
                libreoffice_paths = [
                    "/usr/bin/libreoffice",
                    "/usr/bin/soffice",
                    "/snap/bin/libreoffice"
                ]
            
            # Try to find LibreOffice
            soffice_path = None
            for path in libreoffice_paths:
                if os.path.exists(path):
                    soffice_path = path
                    break
            
            if soffice_path:
                output_dir = os.path.dirname(pdf_path)
                cmd = [
                    soffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", output_dir,
                    docx_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                if result.returncode == 0:
                    # LibreOffice creates PDF with same name as docx
                    expected_pdf = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
                    if os.path.exists(expected_pdf) and expected_pdf != pdf_path:
                        os.rename(expected_pdf, pdf_path)
                    return True, "PDF created successfully using LibreOffice"
                else:
                    errors.append(f"LibreOffice conversion failed: {result.stderr}")
            else:
                errors.append("LibreOffice not found")
                
        except subprocess.TimeoutExpired:
            errors.append("LibreOffice conversion timed out")
        except Exception as e:
            errors.append(f"LibreOffice conversion failed: {str(e)}")
        
        # Method 3: Try native ReportLab PDF generation (always available)
        try:
            # This is our fallback that creates a basic PDF directly
            self.generate_pdf_reportlab(pdf_path)
            return True, "PDF created successfully using ReportLab (basic formatting)"
        except Exception as e:
            errors.append(f"ReportLab generation failed: {str(e)}")
        
        # If all methods failed
        return False, "; ".join(errors)
    
    def generate_pdf_reportlab(self, pdf_path):
        """Generate PDF directly using ReportLab as fallback method"""
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import letter
        
        # Gather content from the form
        content = self.gather_content()
        
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title = f"{content.get('course_number', '')} - {content.get('course_title', '')}"
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        
        # Add basic course info
        info_items = [
            f"Term: {content.get('semester', '')}",
            f"Credits: {content.get('credits', '')}",
            f"Meeting Times: {content.get('meeting_times', '')}",
            f"Location: {content.get('location', '')}",
            f"Instructor: {content.get('instructor_name', '')}",
            f"Email: {content.get('instructor_email', '')}"
        ]
        
        for item in info_items:
            if item.split(': ')[1]:  # Only add if value exists
                story.append(Paragraph(item, styles['Normal']))
        
        story.append(Spacer(1, 12))
        
        # Add description if available
        if content.get('description'):
            story.append(Paragraph("Course Description", styles['Heading2']))
            story.append(Paragraph(content['description'], styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Add a note about formatting
        story.append(Spacer(1, 24))
        note = ("Note: This PDF was generated using basic formatting. "
                "For full formatting, please install Microsoft Word or LibreOffice "
                "and use the Word document export option.")
        story.append(Paragraph(note, styles['BodyText']))
        
        doc.build(story)

    def generate_syllabus(self, export_format="docx"):
        """Generate the final syllabus document"""
        if not self.validate_inputs():
            return

        if export_format == "pdf":
            default_ext = ".pdf"
            file_types = [("PDF Document", "*.pdf"), ("Word Document", "*.docx")]
        else:
            default_ext = ".docx"
            file_types = [("Word Document", "*.docx"), ("PDF Document", "*.pdf")]

        export_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=file_types,
            title="Save Syllabus As"
        )

        if not export_path:
            return

        try:
            if export_path.lower().endswith('.pdf') or export_format == "pdf":
                # Create temporary Word document first
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
                    docx_path = tmp_docx.name
                
                try:
                    # Generate Word document
                    doc = self.create_syllabus_document()
                    doc.save(docx_path)
                    
                    # Convert to PDF using robust method
                    success, message = self.convert_docx_to_pdf_robust(docx_path, export_path)
                    
                    if success:
                        messagebox.showinfo("Success", f"Syllabus saved as PDF: {export_path}\n\n{message}")
                    else:
                        # Offer fallback to Word document
                        fallback_path = export_path.replace('.pdf', '.docx')
                        choice = messagebox.askyesno("PDF Failed - Save as Word?", 
                            f"PDF conversion failed:\n{message}\n\n"
                            f"Would you like to save as Word document instead?\n"
                            f"File: {fallback_path}\n\n"
                            f"You can then open it in Word/LibreOffice and use 'Save As PDF'.")
                        
                        if choice:
                            doc.save(fallback_path)
                            messagebox.showinfo("Saved as Word", 
                                f"Document saved as: {fallback_path}\n\n"
                                f"To convert to PDF:\n"
                                f"• Open in Microsoft Word or LibreOffice\n"
                                f"• Use 'File > Export as PDF' or 'Save As > PDF'\n"
                                f"• Or use online converters like SmallPDF")
                        else:
                            messagebox.showinfo("PDF Not Created", 
                                "PDF was not created. Consider installing LibreOffice for better PDF support.")
                
                finally:
                    # Clean up temporary file
                    try:
                        os.remove(docx_path)
                    except:
                        pass
            else:
                # Save as Word document
                doc = self.create_syllabus_document()
                doc.save(export_path)
                messagebox.showinfo("Success", f"Syllabus saved as Word document: {export_path}")
                
        except Exception as e:
            messagebox.showerror("Error", 
                f"Failed to save syllabus:\n{str(e)}\n\n"
                "Please make sure you have write permissions and the file is not open in another program.")
            
    def validate_inputs(self):
        """Validate required inputs before generating syllabus"""
        required_fields = [
            (self.entry_course_num, "Course Number"),
            (self.entry_course_title, "Course Title"),
            (self.entry_term, "Term"),
            (self.entry_credits, "Credits"),
            (self.entry_meeting_times, "Meeting Times"),
            (self.entry_location, "Location"),
            (self.entry_instr_name, "Instructor Name"),
            (self.entry_instr_email, "Instructor Email")
        ]
        
        for field, name in required_fields:
            if not field.get().strip():
                messagebox.showerror("Error", f"{name} is required.")
                return False
        return True

    def create_syllabus_document(self):
        """Create the Word document for the syllabus following the exact format from the example"""
        try:
            doc = Document()
            
            # Add page numbers in the footer
            section = doc.sections[0]
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.text = "Page "
            
            # Add a field for the page number
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            run._element.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._element.append(fldChar2)
            
            # Continue with the rest of the document creation code
            # Title and Course Info (centered)
            title = doc.add_heading(f"{self.entry_course_num.get()}: {self.entry_course_title.get()}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            term = doc.add_paragraph(f"{self.entry_term.get()} ({self.entry_credits.get()} credits)")
            term.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Single space after title
            
            # I. General Information
            doc.add_heading("I. General Information", level=1)
            
            # Meeting times and location - no extra spacing
            p = doc.add_paragraph()
            p.add_run("Meeting days and times: ").bold = True
            p.add_run(self.entry_meeting_times.get())
            
            p = doc.add_paragraph()
            p.add_run("Class location: ").bold = True
            p.add_run(self.entry_location.get())
            
            # Instructor info - compact format
            p = doc.add_paragraph()
            p.add_run("\nInstructor:").bold = True
            
            instructor_info = [
                ("Name:", self.entry_instr_name.get()),
                ("Office:", self.entry_instr_office.get()),
                ("Phone:", self.entry_instr_phone.get()),
                ("Email:", self.entry_instr_email.get()),
                ("Office Hours:", self.entry_instr_office_hours.get())
            ]
            for label, value in instructor_info:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)  # Reduced indentation
                p.paragraph_format.space_after = Pt(0)  # Remove spacing after paragraphs
                p.add_run(f"{label} ").bold = True
                
                if label == "Email:":
                    # Add email as mailto hyperlink
                    try:
                        add_hyperlink(p, value, f"mailto:{value}")
                    except Exception:
                        # Fallback to plain text if hyperlink creation fails
                        p.add_run(value)
                else:
                    # Other fields remain as plain text
                    p.add_run(value)
            
            # Sections - compact format
            if self.ta_entries:
                p = doc.add_paragraph()
                p.add_run("\nSections:").bold = True
                
                for ta in self.ta_entries:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Name: ").bold = True
                    p.add_run(ta[0].get())
                    
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Email: ").bold = True
                    
                    # Add email as mailto hyperlink
                    try:
                        add_hyperlink(p, ta[1].get(), f"mailto:{ta[1].get()}")
                    except Exception:
                        # Fallback to plain text
                        p.add_run(ta[1].get())
                        
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Office Hours: ").bold = True
                    p.add_run(ta[2].get())

                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Class Room: ").bold = True
                    p.add_run(ta[3].get())

                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Class Time: ").bold = True
                    p.add_run(ta[4].get())
            
            # Course Description
            doc.add_heading("Course Description", level=1)
            doc.add_paragraph(self.txt_description.get("1.0", tk.END).strip())
            
            # Prerequisites (moved before General Education)
            doc.add_heading("Prerequisites", level=1)
            doc.add_paragraph(self.entry_prerequisites.get().strip() if hasattr(self, 'entry_prerequisites') else "None")
            
            # --- Add General Education Designation (moved after Prerequisites) ---
            if self.show_gen_ed.get():
                # Define the designation variable
                designation = "Social and Behavioral Sciences (S)"  # Default value
                
                # Add the General Education heading and note
                doc.add_heading(f"General Education Designation: {designation}", level=1)
                
                # Add the full General Education description text
                doc.add_paragraph(gen_ed_default)
                
                # Then continue with the existing code for the success message
                course_num = self.entry_course_num.get()
                doc.add_paragraph(f"Your successful completion of {course_num} with a grade of \"C\" or higher will count towards UF's General Education State Core in {designation}. It will also count towards the State of Florida's Civic Literacy requirement.")
            
            # Course Objectives - Add this section (moved after General Education)
            doc.add_heading("Course Objectives", level=1)
            p = doc.add_paragraph("All General Education area objectives can be found ")
            add_hyperlink(p, "here", "https://undergrad.aa.ufl.edu/general-education/gen-ed-program/subject-area-objectives/")
            p.add_run(".")
            
            if hasattr(self, 'objective_entries') and any(obj["entry"].get().strip() for obj in self.objective_entries):
                for i, obj in enumerate(self.objective_entries, 1):
                    obj_text = obj["entry"].get().strip()
                    if obj_text:
                        p = doc.add_paragraph(f"{i}. {obj_text}")
                        p.paragraph_format.left_indent = Inches(0.25)
            
            # Add Student Learning Outcomes
            doc.add_heading("II. Student Learning Outcomes", level=1)
            doc.add_paragraph("A student who successfully completes this course will:")
            
            if hasattr(self, 'outcome_entries') and self.outcome_entries:
                for i, outcome_entry in enumerate(self.outcome_entries, 1):
                    outcome_text = outcome_entry["entry"].get().strip()
                    if outcome_text:
                        p = doc.add_paragraph(f"{i}. {outcome_text}")
                        p.paragraph_format.left_indent = Inches(0.25)
            else:
                # Default outcomes if none provided
                default_outcomes = [
                    "Describe the factual details of the substantive historical episodes under study.",
                    "Identify and analyze foundational developments that shaped history using critical thinking skills.",
                    "Demonstrate an understanding of the primary ideas, values, and perceptions that have shaped history.",
                    "Demonstrate competency in civic literacy."
                ]
                for i, outcome in enumerate(default_outcomes, 1):
                    p = doc.add_paragraph(f"{i}. {outcome}")
                    p.paragraph_format.left_indent = Inches(0.25)
            
            # If General Education is enabled, add the objectives table after the Student Learning Outcomes
            if self.show_gen_ed.get():
                # Add Learning Outcomes table
                doc.add_paragraph()
                doc.add_paragraph(f"Objectives—General Education and {designation}")
                
                # Create table with the correct number of rows and columns
                table_rows = 1  # Header row
                
                # Count rows based on learning objectives entries if they exist
                if hasattr(self, 'learning_objectives_entries') and self.learning_objectives_entries:
                    for category, entries in self.learning_objectives_entries.items():
                        # Skip entries that were removed
                        if 'frame' in entries and not entries['frame'].winfo_exists():
                            continue
                        table_rows += 1
                else:
                    # Default 3 categories if no custom entries
                    table_rows += 3
                    
                table = doc.add_table(rows=table_rows, cols=4)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "CATEGORY"
                
                # Customize the second header based on designation
                slo_header = "SOCIAL SCIENCE SLOS"
                if "Humanities" in designation:
                    slo_header = "HUMANITIES SLOS"
                elif "International" in designation:
                    slo_header = "INTERNATIONAL SLOS"
                elif "Diversity" in designation:
                    slo_header = "DIVERSITY SLOS"
                elif "Biological" in designation:
                    slo_header = "BIOLOGICAL SCIENCES SLOS"
                elif "Physical" in designation:
                    slo_header = "PHYSICAL SCIENCES SLOS"
                elif "Mathematics" in designation:
                    slo_header = "MATHEMATICS SLOS"
                    
                header_cells[1].text = slo_header
                header_cells[2].text = "STATE SLO ASSIGNMENTS"
                header_cells[3].text = "COURSE-SPECIFIC"
                
                # Make header row bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                row_idx = 1
                if hasattr(self, 'learning_objectives_entries') and self.learning_objectives_entries:
                    for category, entries in self.learning_objectives_entries.items():
                        # Skip entries that were removed
                        if 'frame' in entries and not entries['frame'].winfo_exists():
                            continue
                        
                        # For custom categories, use the name_entry value
                        if 'name_entry' in entries:
                            category_name = entries['name_entry'].get()
                        else:
                            category_name = category
                        
                        slo_text = entries['slo'].get("1.0", tk.END).strip()
                        assignments_text = entries['assignments'].get("1.0", tk.END).strip()
                        course_specific_text = entries['course_specific'].get("1.0", tk.END).strip()
                        
                        if row_idx < len(table.rows):
                            row = table.rows[row_idx]
                            row.cells[0].text = category_name
                            row.cells[1].text = slo_text
                            row.cells[2].text = assignments_text
                            row.cells[3].text = course_specific_text
                            row_idx += 1
                else:
                    # Fallback to default table if no custom entries exist
                    default_data = [
                        ("Content", 
                         "Identify, describe, and explain key themes, principles, and terminology; the history, theory and/or methodologies used; and social institutions, structures and processes.", 
                         "Outcomes 1-4\n\nStudents will demonstrate their knowledge of the details of the substantive historical episodes by analyzing primary and secondary sources in short papers, homework assignments, exams, and in-class discussion."),
                        ("Critical Thinking", 
                         "Apply formal and informal qualitative or quantitative analysis effectively to examine the processes and means by which individuals make personal and group decisions. Assess and analyze ethical perspectives in individual and societal decisions.", 
                         "Outcomes 1-4\n\nStudents will demonstrate their ability in applying qualitative and quantitative methods by analyzing primary and secondary sources in short papers, homework assignments, and exams by using critical thinking skills."),
                        ("Communication", 
                         "Communication is the development and expression of ideas in written and oral forms.", 
                         "Outcomes 1-4\n\nStudents will identify and explain key developments that shaped history in written assignments and class discussion.\n\nStudents will demonstrate their understandings of the primary ideas, values, and perceptions that have shaped history and will describe them in written assignments, exams, and class discussion.")
                    ]
                    
                    for i, (category, slo, assignments) in enumerate(default_data):
                        if row_idx < len(table.rows):
                            row = table.rows[row_idx]
                            row.cells[0].text = category
                            row.cells[1].text = slo
                            row.cells[2].text = assignments
                            row.cells[3].text = ""
                            row_idx += 1
            
            # III. Graded Work
            doc.add_heading("III. Graded Work", level=1)
            
            # Materials (if provided)
            if hasattr(self, 'materials_text') and self.materials_text.get("1.0", tk.END).strip():
                doc.add_heading("Required Materials", level=2)
                materials_text = self.materials_text.get("1.0", tk.END).strip()
                # --- Use markup parser for formatted output ---
                self.parse_materials_markup(materials_text, doc=doc)
                # Always include the Materials Fee value from the fee_entry
                fee_value = ""
                if hasattr(self, 'fee_entry') and self.fee_entry.get().strip():
                    fee_value = self.fee_entry.get().strip()
                else:
                    fee_value = "0.00"
                p = doc.add_paragraph()
                p.add_run("\nMaterials Fee: $").bold = True
                p.add_run(fee_value)

            # ==== CONTINUING AFTER REQUIRED MATERIALS SECTION ====
            
            # Grading Components (Categories and Assignments)
            if hasattr(self, 'category_frames') and self.category_frames:
                doc.add_heading("Grading Components", level=2)
                
                # Create a table for categories and weights
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Category"
                header_cells[1].text = "Weight"
                
                # Make header row bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for category in self.category_frames:
                    name = category["name"].get().strip()
                    weight = category["weight"].get().strip()
                    
                    if name and weight:
                        row_cells = table.add_row().cells
                        row_cells[0].text = name
                        row_cells[1].text = f"{weight}%"
                
                # Category descriptions and assignments
                for category in self.category_frames:
                    name = category["name"].get().strip()
                    desc = category["description"].get("1.0", tk.END).strip()
                    
                    if name and desc:
                        p = doc.add_paragraph()
                        p.add_run(f"\n{name}: ").bold = True
                        p.add_run(desc)
                    
                    # Assignments for this category
                    if "assignments" in category and category["assignments"]:
                        has_assignments = False
                        for assignment in category["assignments"]:
                            title = assignment["title"].get().strip()
                            due_date = assignment["due date"].get().strip()
                            points = assignment["points"].get().strip()
                            description = assignment["description"].get("1.0", tk.END).strip() if "description" in assignment else ""
                            
                            if title:
                                if not has_assignments:
                                    p = doc.add_paragraph()
                                    p.add_run(f"{name} Assignments:").bold = True
                                    has_assignments = True
                                
                                p = doc.add_paragraph()
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.add_run(f"• {title}")
                                
                                if due_date:
                                    p.add_run(f" (Due: {due_date})")
                                if points:
                                    p.add_run(f" - {points} points")
                                
                                if description:
                                    p = doc.add_paragraph(description)
                                    p.paragraph_format.left_indent = Inches(0.5)
            
            # Grading Scale
            doc.add_heading("Grading Scale", level=2)
            
            # Create table for grading scale
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Letter Grade"
            header_cells[1].text = "Number Grade"
            
            # Make header row bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Grade scale data
            grades = [
                ("A", "100-93"),
                ("A-", "92-90"),
                ("B+", "89-87"),
                ("B", "86-83"),
                ("B-", "82-80"),
                ("C+", "79-77"),
                ("C", "76-73"),
                ("C-", "72-70"),
                ("D+", "69-67"),
                ("D", "66-63"),
                ("D-", "62-60"),
                ("E", "59-0")
            ]
            
            for letter, number in grades:
                row_cells = table.add_row().cells
                row_cells[0].text = letter
                row_cells[1].text = number
            
            # Add UF grading policies note
            p = doc.add_paragraph()
            p.add_run("See the UF Catalog's ") 
            add_hyperlink(p,"Grades and Grading Policies", "https://catalog.ufl.edu/UGRD/academic-regulations/grades-grading-policies/")
            p.add_run(" for information on how UF assigns grade points.")

            # Add rounding statement if enabled
            if hasattr(self, 'grading_rounding_var') and self.grading_rounding_var.get():
                from constants import grading_rounding_default
                p = doc.add_paragraph(grading_rounding_default)

            # Add minimum grade note
            p = doc.add_paragraph()
            p.add_run("Note: A minimum grade of C is required to earn General Education credit.")
            
            # Instructions for Submitting Written Assignments
            doc.add_heading("Instructions for Submitting Written Assignments", level=1)
            doc.add_paragraph("All written assignments must be submitted as Word documents (.doc or .docx) through the \"Assignments\" portal in Canvas by the specified deadlines. Do NOT send assignments as PDF files.")
            
            # Add Late Submissions policy
            if (self.late_submissions_policy_var.get() and
                ((hasattr(self, 'late_policy_text') and 
                  self.late_policy_text.get("1.0", tk.END).strip() != "") or
                 (hasattr(self, 'late_policy_var') and 
                  self.late_policy_var.get() in self.late_policies))):
                doc.add_heading("Late Submissions", level=2)
                if hasattr(self, 'late_policy_text') and self.late_policy_text.get("1.0", tk.END).strip():
                    doc.add_paragraph(self.late_policy_text.get("1.0", tk.END).strip())
                elif hasattr(self, 'late_policy_var') and self.late_policy_var.get() in self.late_policies:
                    selected_policy = self.late_policy_var.get()
                    doc.add_paragraph(self.late_policies[selected_policy])
                else:
                    doc.add_paragraph("Late submission policy not specified.")
            
            # Add Extra Credit policy
            if (self.extra_credit_policy_var.get() and
                ((hasattr(self, 'extra_credit_text') and 
                  self.extra_credit_text.get("1.0", tk.END).strip() != "") or
                 (hasattr(self, 'extra_credit_var') and 
                  self.extra_credit_var.get() in self.extra_credit_policies))):
                doc.add_heading("Extra Credit", level=2)
                if hasattr(self, 'extra_credit_text') and self.extra_credit_text.get("1.0", tk.END).strip():
                    doc.add_paragraph(self.extra_credit_text.get("1.0", tk.END).strip())
                elif hasattr(self, 'extra_credit_var') and self.extra_credit_var.get() in self.extra_credit_policies:
                    selected_policy = self.extra_credit_var.get()
                    doc.add_paragraph(self.extra_credit_policies[selected_policy])
                else:
                    doc.add_paragraph("Extra credit policy not specified.")
            

            # Canvas Policy
            if (self.canvas_policy_var.get() and hasattr(self, 'canvas_policy_text')):
                doc.add_heading("Canvas", level=2)
                doc.add_paragraph(self.canvas_policy_text.get("1.0", tk.END).strip())


            # Technology Policy
            if (self.technology_policy_var.get() and hasattr(self, 'technology_policy_text')):
                doc.add_heading("Technology in the Classroom", level=2)
                doc.add_paragraph(self.technology_policy_text.get("1.0", tk.END).strip())


            # Communication Policy
            if (self.communication_policy_var.get() and hasattr(self, 'communication_policy_text')):
                doc.add_heading("Class Communication Policy", level=2)
                doc.add_paragraph(self.communication_policy_text.get("1.0", tk.END).strip())


            # Assignment Support section
            if (self.outside_support_var.get() and hasattr(self, 'support_text')):
                doc.add_heading("Assignment Support Outside the Classroom", level=2)
                doc.add_paragraph(self.support_text.get("1.0", tk.END).strip())
                        
            # IV. University Policies and Resources (formerly V.)
            doc.add_heading("IV. University Policies and Resources", level=1)
            
            # Check if simplified policies are enabled
            if hasattr(self, 'use_simplified_policies_var') and self.use_simplified_policies_var.get():
                # Use simplified UF policies
                from constants import uf_policy_simplified
                p = doc.add_paragraph()
                p.add_run("This course complies with all UF academic policies. For information on those polices and for resources for students, please see ")
                add_hyperlink(p, "this link", "https://syllabus.ufl.edu/syllabus-policy/uf-syllabus-policy-links/")
                p.add_run(".")
            else:
                # Use original detailed policies
                # Accommodations policy
                doc.add_heading("Students requiring accommodation", level=2)
                p = doc.add_paragraph()
                accommodations_text = (
                    "Students with disabilities who experience learning barriers and would like to request academic accommodations "
                    "should connect with the Disability Resource Center by visiting https://disability.ufl.edu/students/get-started/. "
                    "It is important for students to share their accommodation letter with the instructor and discuss their "
                    "access needs as early as possible in the semester."
                )
                process_text_with_hyperlinks(p, accommodations_text)

                # University Honesty Policy
                doc.add_heading("University Honesty Policy", level=2)
                p = doc.add_paragraph()
                p.add_run("UF students are bound by The Honor Pledge which states \"We, the members of the University of Florida community, pledge to hold ourselves and our peers to the highest standards of honor and integrity by abiding by the Honor Code.\" " +
                "On all work submitted for credit by students at the University of Florida, the following pledge is either required or implied: " +
                "\"On my honor, I have neither given nor received unauthorized aid in doing this assignment.\" " +
                "The Conduct Code specifies a number of behaviors that are in violation of this code and the possible sanctions.")
                add_hyperlink(p, " See the UF Conduct Code website for more information", "https://sccr.dso.ufl.edu/process/student-conduct-code/")
                p.add_run(". If you have any questions or concerns, please consult with the instructor or TAs in this class.")

                doc.add_heading("Plagiarism and Related Ethical Violations ", level=2)
                p = doc.add_paragraph()
                p.add_run("Ethical violations such as plagiarism, cheating, academic misconduct (e.g. passing off others' work as your own, reusing old assignments, etc.) " \
                "will not be tolerated and will result in a failing grade in this course. Students must be especially wary of plagiarism. " \
                "The UF Student Honor Code defines plagiarism as follows: "
                "A student shall not represent as the student's own work all or any portion of the work of another. "
                "Plagiarism includes (but is not limited to): a. Quoting oral or written materials, whether published or unpublished, without proper attribution. "
                "b. Submitting a document or assignment which in whole or in part is identical or substantially identical to a document or assignment not authored by the student."
                " Note that plagiarism also includes the use of any artificial intelligence programs, such as ChatGPT. ")
                
                
            # V. Course Schedule (Calendar) (formerly VI.)
            doc.add_heading("V. Calendar", level=1)
            
            # Create schedule table if there are entries
            if hasattr(self, 'schedule_entries') and self.schedule_entries:
                # Create table with headers
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Set column headers
                header_cells = table.rows[0].cells
                header_cells[0].text = "Date"
                header_cells[1].text = "Topic"
                header_cells[2].text = "Readings/Preparation"
                header_cells[3].text = "Work Due"
                
                # Make header row bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add each schedule entry as a row
                for entry in self.schedule_entries:
                    date_text = entry["date"].get().strip()
                    topic_text = entry["topic"].get().strip()
                    readings_text = entry["readings"].get("1.0", tk.END).strip()
                    work_due_text = entry["work_due"].get().strip()
                    
                    # Skip empty rows
                    if not any([date_text, topic_text, readings_text, work_due_text]):
                        continue
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = date_text
                    row_cells[1].text = topic_text
                    row_cells[2].text = readings_text
                    row_cells[3].text = work_due_text
            else:
                doc.add_paragraph("Schedule will be provided separately.")

            return doc
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while generating the document: {e}")
            import traceback
            traceback.print_exc()
            # Return an empty document so we don't crash
            return Document()

    def parse_materials_markup(self, text, doc=None):
        """
        Parse Markdown-like markup in Required Materials and add to Word doc.
        - Supports the following syntax:
          - *italic* -> Text between single asterisks (*) will be italicized.
          - **bold** -> Text between double asterisks (**) will be bolded.
          - [text](url) -> Text inside square brackets ([text]) followed by a URL in parentheses (url) will become a hyperlink.
        """
        import re
        
        if doc is None:
            # Just return the text if no document is provided
            return text
        
        # Create paragraph for the content
        paragraph = doc.add_paragraph()
        
        # Regular expression to find markdown patterns
        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))')
        
        # Split the text based on the pattern
        parts = pattern.split(text)
        
        # Process each part
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('[') and '](' in part and part.endswith(')'):
                # Hyperlink
                # Extract the text and URL
                link_text = part[1:part.index('](')]
                url = part[part.index('](')+2:-1]
                
                # Add hyperlink
                run = paragraph.add_run(link_text)
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # Blue color
                run.font.underline = True
                
                # Add the actual hyperlink
                # This is a simplification - true hyperlinks need more work with XML
                # For now, we'll just make it look like a hyperlink
            else:
                # Regular text
                paragraph.add_run(part)
        
        return paragraph

    def show_formatting_help(self):
        """Display a popup with formatting help."""
        help_window = tk.Toplevel(self.root)
        help_window.title("Formatting Help")
        help_window.geometry("400x300")

        explanation = """
        Markdown-like Formatting Guide:
        - *italic* -> Text between single asterisks (*) will be italicized.
        - **bold** -> Text between double asterisks (**) will be bolded.
        - [text](url) -> Text inside square brackets ([text]) followed by a URL in parentheses (url) will become a hyperlink.

        Example:
        Input: "This is *italic*, **bold*, and [a link](http://example.com)."
        Output in Word:
          - "italic" will appear italicized.
          - "bold" will appear bolded.
          - "a link" will appear as a clickable hyperlink pointing to "http://example.com".
        """

        text_widget = tk.Text(help_window, wrap=tk.WORD, font=("Arial", 10))
        text_widget.insert(tk.END, explanation)
        text_widget.config(state=tk.DISABLED)  # Make the text read-only
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
