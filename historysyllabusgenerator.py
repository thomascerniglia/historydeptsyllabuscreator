import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from tkinter import scrolledtext
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from tkinter import font as tkfont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Spacer
import json
import os
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pickle
import sys
import re
import tempfile
import docx
try:
    # Import the convert function from docx2pdf
    from docx2pdf import convert
    
except ImportError:
    # Handle case when docx2pdf is not installed
    messagebox.showwarning("Missing Library", "The docx2pdf library is not installed. PDF export will not work.")
    # Define a dummy function to prevent errors
    def convert(input_path, output_path):
        raise RuntimeError("docx2pdf library is not installed. Cannot convert to PDF.")
    
# Default template text for various sections
course_description_default = ""
prerequisites_default = "None"

gen_ed_default = (
    "Social Science courses must afford students an understanding of the basic social and behavioral "
    "science concepts and principles used in the analysis of behavior and past and present social, "
    "political, and economic issues. Social and Behavioral Sciences (S) is a sub-designation of Social Sciences "
    "at the University of Florida. These courses provide instruction in the history, key themes, principles, "
    "terminology, and underlying theory or methodologies used in the social and behavioral sciences. Students "
    "will learn to identify, describe and explain social institutions, structures or processes. These courses "
    "emphasize the effective application of accepted problem-solving techniques. Students will apply formal "
    "and informal qualitative or quantitative analysis to examine the processes and means by which individuals "
    "make personal and group decisions, as well as the evaluation of opinions, outcomes or human behavior. "
    "Students are expected to assess and analyze ethical perspectives in individual and societal decisions.\n\n"
    
)

honesty_plagiarism_default = (
    "UF students are bound by The Honor Pledge which states: \"We, the members of the University of "
    "Florida community, pledge to hold ourselves and our peers to the highest standards of honor and "
    "integrity by abiding by the Honor Code. On all work submitted for credit by students at the "
    "University of Florida, the following pledge is either required or implied: 'On my honor, I have "
    "neither given nor received unauthorized aid in doing this assignment.' The Conduct Code specifies "
    "a number of behaviors that are in violation of this code and the possible sanctions. If you have "
    "any questions or concerns, please consult with the instructor or TAs in this class.\n\n"
    "Ethical violations such as plagiarism, cheating, and other academic misconduct will not be "
    "tolerated and will result in a failing grade in this course. Note that plagiarism also includes "
    "the use of any artificial intelligence program (e.g., ChatGPT) to produce work for this course."
)

recording_policy_default = (
    "Students are allowed to record video or audio of class lectures. However, the purposes for "
    "which these recordings may be used are strictly controlled. The only allowable purposes are "
    "(1) for personal educational use, (2) in connection with a complaint to the university, or "
    "(3) as evidence in, or in preparation for, a criminal or civil proceeding. All other purposes "
    "are prohibited. Specifically, students may not publish recorded lectures without the written "
    "consent of the instructor.\n\n"
    "A \"class lecture\" is an educational presentation intended to inform or teach enrolled students "
    "about a particular subject, including any instructor-led discussions that form part of the "
    "presentation, and delivered by any instructor hired or appointed by the University, or by a "
    "guest instructor, as part of a University of Florida course.\n\n"
    "Publication without permission of the instructor is prohibited. To \"publish\" means to share, "
    "transmit, circulate, distribute, or provide access to a recording, regardless of format or "
    "medium, to another person (or persons), including but not limited to another student within "
    "the same class section. A student who publishes a recording without written consent may be "
    "subject to a civil cause of action instituted by a person injured by the publication and/or "
    "discipline under UF Regulation 4.040 Student Honor Code and Student Conduct Code."
)

accommodations_default = (
    "Students with disabilities who experience learning barriers and would like to request academic accommodations should connect with the Disability Resource Center by visiting **https://disability.ufl.edu/students/get-started/**. "
    "It is important for students to share their accommodation letter with the instructor and discuss their access needs as early as possible in the semester."
)

canvas_policy_default = (
    "-Replace with your Canvas Policies-"
)

technology_policy_default = (
    "-Replace with your Technology Policies-"
)

evaluations_default = (
    "Students are expected to provide professional and respectful feedback on the quality of instruction by completing **course evaluations** online via GatorEvals. "
    "Evaluations can be done via the email link from GatorEvals, the link in Canvas, or by logging in to the GatorEvals portal. Students will be notified when the evaluation period opens, and can view summary results of past evaluations on the GatorEvals website."
)

conflict_resolution_default = (
    "Any classroom issues, disagreements, or grade disputes should first be discussed between the student and instructor. "
    "If the problem cannot be resolved, please contact the Associate Chair of the History Department. Be prepared to provide documentation of the issue and all graded materials. "
    "Unresolved issues may be referred to the University Ombuds Office or the Dean of Students Office for further resolution."
)

campus_resources_default = """U Matter, We Care: If you or someone you know is in distress, please contact umatter@ufl.edu, 352-392-1575, or visit U Matter, We Care website to refer or report a concern and a team member will reach out to the student in distress.

Counseling and Wellness Center: Visit the Counseling and Wellness Center website or call 352-392-1575 for information on crisis services as well as non-crisis services.

Student Health Care Center: Call 352-392-1161 for 24/7 information to help you find the care you need, or visit the Student Health Care Center website.

University Police Department: Visit UF Police Department website or call 352-392-1111 (or 9-1-1 for emergencies).

UF Health Shands Emergency Room / Trauma Center: For immediate medical care call 352-733-0111 or go to the emergency room at 1515 SW Archer Road, Gainesville, FL 32608; Visit the UF Health Emergency Room and Trauma Center website.

GatorWell Health Promotion Services: For prevention services focused on optimal wellbeing, including Wellness Coaching for Academic Success, visit the GatorWell website or call 352-273-4450.

Student Success Initiative, https://studentsuccess.ufl.edu/.

Field and Fork Pantry. Food and toiletries for students experiencing food insecurity.  

Dean of Students Office. 202 Peabody Hall, 392-1261. Among other services, the DSO assists students who are experiencing situations that compromises their ability to attend classes. This includes family emergencies and medical issues (including mental health crises)."""

# Define default text for Academic Resources
academic_resources_default = """E-learning technical support: Contact the UF Computing Help Desk at 352-392-4357 or via e-mail at helpdesk@ufl.edu.  

Career Connections Center: Reitz Union Suite 1300, 352-392-1601. Career assistance and counseling services.  

Library Support: Various ways to receive assistance with respect to using the libraries or finding resources.  

Teaching Center: Broward Hall, 352-392-2010 or to make an appointment 352- 392-6420. General study skills and tutoring.  

Writing Studio: 2215 Turlington Hall, 352-846-1138. Help brainstorming, formatting, and writing papers.  

Student Complaints On-Campus: Visit the Student Honor Code and Student Conduct Code webpage for more information.  

On-Line Students Complaints: View the Distance Learning Student Complaint Process."""

class SyllabusTemplate:
    """Class to represent a syllabus template"""
    def __init__(self, course_code, title, description="", objectives=None, outcomes=None):
        self.course_code = course_code
        self.title = title
        self.description = description
        self.objectives = objectives or []
        self.outcomes = outcomes or []
        self.schedule = []
        self.grading_categories = []
        self.learning_objectives = {}

class HistorySyllabusGenerator:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("History Syllabus Generator")
        self.root.state('zoomed')
        self.current_template = None
        self.templates = []
        self.template_names = []
        # Store TAs
        self.ta_entries = []
        # Store schedule entries
        self.schedule_entries = []
        # Store category frames
        self.category_frames = []
        # Store learning objective entries
        self.learning_objectives_entries = {}
        # Store outcome entries
        self.outcome_entries = []
        # Set up styles
        self.setup_styles()
        # Add variable for Gen Ed toggle
        self.show_gen_ed = tk.BooleanVar(value=True)
        self.outside_support_var = tk.BooleanVar(value=True)
        self.in_class_recording_var = tk.BooleanVar(value=True)
        self.conflict_resolution_var = tk.BooleanVar(value=True)
        self.campus_resources_var = tk.BooleanVar(value=True)
        self.academic_resources_var = tk.BooleanVar(value=True)
        self.late_work_var = tk.BooleanVar(value=True)
        self.optional_policies = {
            "in_class_recording": self.in_class_recording_var,
            "conflict_resolution": self.conflict_resolution_var,
            "campus_resources": self.campus_resources_var,
            "academic_resources": self.academic_resources_var,
            "late_work": self.late_work_var
        }
        self.create_main_interface()  # <-- Only call here!
        self.templates = self.load_default_templates()
        
        self.template_names = [f"{t.course_code}: {t.title}" for t in self.templates]
        self.template_combo['values'] = ["Clear Template"] + self.template_names

    def load_default_templates(self):
        """Load a single default template for testing purposes"""
        try:
            # Create and initialize a test template
            test_template = SyllabusTemplate(
                course_code="TEST1000",
                title="TEST TEMPLATE",
                description="This is a test template for the History Syllabus Generator.",
                objectives=[
                    "Understand the principles of historical research.",
                    "Develop critical thinking skills through analysis of primary sources.",
                    "Learn to construct historical arguments based on evidence."
                ],
                outcomes=[
                    "Demonstrate ability to analyze primary sources.",
                    "Write clear and coherent historical essays.",
                    "Present research findings effectively."
                ]
            )
            # Then set all the properties
            test_template.instructor_name = "Dr. Jane Doe"
            test_template.prerequisites = "HIS100 or instructor permission" 
            test_template.instructor_office = "Room 123, History Building"
            test_template.instructor_phone = "555-123-4567"
            test_template.instructor_email = "jane.doe@university.edu"
            test_template.instructor_office_hours = "MWF 2:00 PM - 4:00 PM" 
            test_template.semester = "Spring 2025"
            test_template.credits = "3"
            test_template.class_days = "MWF"
            test_template.class_times = "10:00 AM - 10:50 AM"
            test_template.classroom = "History Building 202"
            # Add TAs with complete information
            test_template.tas = [
                {
                    "name": "John Smith",
                    "email": "john.smith@university.edu",
                    "office_hours": "TTh 10:00 AM - 12:00 PM"
                },
                {
                    "name": "Emily Johnson",
                    "email": "emily.johnson@university.edu",
                    "office_hours": "MW 1:00 PM - 3:00 PM"
                }
            ]
            test_template.learning_objectives = {
                "Content": {
                    "slo": "Identify, describe, and explain key themes, principles, and terminology; the history, theory and/or methodologies used; and social institutions, structures and processes.",
                    "assignments": "Outcomes 1-4",
                    "course_specific": "Students will analyze primary and secondary sources in short papers, homework assignments, exams, and in-class discussion."
                },
                "Critical Thinking": {
                    "slo": "Apply formal and informal qualitative or quantitative analysis effectively to examine the processes and means by which individuals make personal and group decisions. Assess and analyze ethical perspectives in individual and societal decisions.",
                    "assignments": "Outcomes 1-4",
                    "course_specific": "Students will apply critical thinking skills in written assignments and exams."
                },
                "Communication": {
                    "slo": "Communication is the development and expression of ideas in written and oral forms.",
                    "assignments": "Outcomes 1-4",
                    "course_specific": "Students will present research findings and participate in class discussions."
                }
            }
            # Set default choices for late submissions and extra credit policy dropdowns
            test_template.late_policy = "Standard (10% per day)"
            test_template.extra_credit_policy = "Standard"

            # --- Add sample schedule entries for testing ---
            test_template.schedule = [
                {
                    "date": "January 13, 2025",
                    "topic": "Syllabus Review; Reconstruction",
                    "readings": "AMH 2020 Syllabus [825 words]\n'Reconstruction,' Chapter 15, American Yawp [10390 words]",
                    "work_due": "Syllabus Quiz due by 11:59pm"
                },
                {
                    "date": "January 15, 2025",
                    "topic": "Reconstruction",
                    "readings": "Frederick Douglass, 'Remembering the Civil War' (1878)\npp. canonsociety.org/the-civil-war-1867 [1006 words]",
                    "work_due": "Reading Response #1"
                },
                {
                    "date": "January 17, 2025",
                    "topic": "TA Session #1",
                    "readings": "All January 15 Readings",
                    "work_due": "Discussion Board Post"
                },
                {
                    "date": "January 20, 2025",
                    "topic": "No Class (Holiday)",
                    "readings": "No readings assigned",
                    "work_due": "None"
                },
                {
                    "date": "January 22, 2025",
                    "topic": "The New South",
                    "readings": "Henry Grady, 'The New South' Speech (1886)\nAmerican Yawp, Chapter 16 excerpt",
                    "work_due": "Short Essay #1 due"
                },
                {
                    "date": "January 24, 2025",
                    "topic": "TA Session #2",
                    "readings": "All January 22 Readings",
                    "work_due": "Quiz #1"
                },
                {
                    "date": "January 27, 2025",
                    "topic": "Gilded Age Politics",
                    "readings": "American Yawp, Chapter 18\nSelections from Nast Cartoons",
                    "work_due": "Reading Response #2"
                },
                {
                    "date": "January 29, 2025",
                    "topic": "Labor in the Gilded Age",
                    "readings": "Jacob Riis, 'The Working Girls of New York'\nAmerican Yawp, Chapter 18 (cont.)",
                    "work_due": "Short Essay #2 due"
                }
            ]
            # --- Add AMH2020 template ---
            amh2020_template = SyllabusTemplate(
                course_code="AMH2020",
                title="United States Since 1877",
                description=(
                    "In this course, students will trace the history of the United States from the end of the Reconstruction era to the contemporary era. "
                    "Topics will include but are not limited to the rise of Industrialization, the United States’ emergence as an actor on the world stage, "
                    "Constitutional amendments and their impact, the Progressive era, World War I, the Great Depression and New Deal, World War II, the Civil Rights era, "
                    "the Cold War, and the United States since 1989.\n\n"
                    "NOTE: All topics in this course will be taught objectively as objects of analysis, without endorsement of particular viewpoints, and will be observed from multiple perspectives. "
                    "No lesson is intended to espouse, promote, advance, inculcate, or compel a particular feeling, perception, or belief. Students are encouraged to employ critical thinking and to rely on data and verifiable sources to explore readings and subject matter in this course. All perspectives will be respected in class discussions."
                ),
                objectives=[
                    "Address how the Civil War and Reconstruction set the stage for the development of the modern United States.",
                    "Explore how US involvement in the Spanish-American War, World War One, and World War Two reshaped US foreign policy and civil society.",
                    "Present the origins of the Cold War, its implications for US international relations, and its influence on American political culture.",
                    "Enable students to analyze and evaluate the origins and influences of the civil rights movement, the Vietnam War, the women’s movement, and New Right conservatism.",
                    "Teach students how to analyze historical documents and scholarship from a range of authors and time periods."
                ],
                outcomes=[
                    "Describe the factual details of the substantive historical episodes under study.",
                    "Identify and analyze foundational developments that shaped American history since 1877 using critical thinking skills.",
                    "Demonstrate an understanding of the primary ideas, values, and perceptions that have shaped American history.",
                    "Demonstrate competency in civic literacy."
                ]
            )
            amh2020_template.prerequisites = "None."
            amh2020_template.semester = "Spring 2025"
            amh2020_template.credits = "3"
            amh2020_template.class_days = "M, W"
            amh2020_template.class_times = "12:50p - 1:40p"
            amh2020_template.classroom = "MCCC 0100"
            # Instructor/TA fields left blank for user to fill in
            amh2020_template.tas = []
            amh2020_template.learning_objectives = {
                "Content": {
                    "slo": "Identify, describe, and explain key themes, principles, and terminology; the history, theory and/or methodologies used; and social institutions, structures and processes.",
                    "assignments": "Outcomes 1-4",
                    "course_specific": "Students will demonstrate their knowledge of the details of the substantive historical episodes of US History since 1877 by analyzing primary and secondary sources in short papers, homework assignments, exams, and in-class discussion."
                },
                "Critical Thinking": {
                    "slo": "Apply formal and informal qualitative or quantitative analysis effectively to examine the processes and means by which individuals make personal and group decisions. Assess and analyze ethical perspectives in individual and societal decisions.",
                    "assignments": "Outcomes 1-4",
                    "course_specific": "Students will demonstrate their ability in applying qualitative and quantitative methods by analyzing primary and secondary sources in short papers, homework assignments, and exams by using critical thinking skills."
                },
                "Communication": {
                    "slo": "Communication is the development and expression of ideas in written and oral forms.",
                    "assignments": "Outcomes 1-4",
                    "course_specific": (
                        "Students will identify and explain key developments that shaped United States history since 1877 in written assignments and class discussion.\n\n"
                        "Students will demonstrate their understandings of the primary ideas, values, and perceptions that have shaped United States history and will describe them in written assignments, exams, and class discussion."
                    )
                }
            }
            amh2020_template.late_policy = "Custom"
            amh2020_template.extra_credit_policy = "Custom"
            amh2020_template.grading_categories = [
                {
                    "name": "Discussion Section Attendance",
                    "weight": "10",
                    "description": (
                        "Your TA will maintain attendance records for all discussion sections. Students are permitted one unexcused absence without penalty, though they are still required to complete all discussion section homework. "
                        "After the first unexcused absence, each subsequent unexcused absence will result in a one-point deduction from your overall course grade. Every two late arrivals to a discussion section counts as one unexcused absence. "
                        "Students will not be penalized for university-excused absences; see UF’s excused absence policy."
                    ),
                    "assignments": []
                },
                {
                    "name": "Discussion Section Homework and Participation",
                    "weight": "15",
                    "description": (
                        "Participation is assessed based on section homework and meaningful contributions to class discussion. Homework will be announced on Monday and must be submitted via Canvas by 9:30a on Friday. "
                        "This grade will be assessed by your TA, who will assign a participation score of 1-15 using the following rubric:\n\n"
                        "1. (5 points) Student completes the five homework assignments (1 point/homework).\n"
                        "2. (7 points) Student attends ten or more sections (.7 points/section attendance)\n"
                        "3. (3 points) Student contributes actively to ten section meetings (.3 points/contribution)"
                    ),
                    "assignments": []
                },
                {
                    "name": "Midterm Exam",
                    "weight": "20",
                    "description": (
                        "This timed, closed-book exam will draw from the course's lectures, discussions, and readings. A study guide will be posted to Canvas."
                    ),
                    "assignments": [
                        {
                            "title": "Midterm Exam",
                            "due_date": "",
                            "points": "",
                            "description": "Exam date and details to be announced in class and on Canvas."
                        }
                    ]
                },
                {
                    "name": "Final Exam",
                    "weight": "30",
                    "description": (
                        "This take-home exam will draw from the course's lectures, discussions, and readings. A study guide will be posted to Canvas."
                    ),
                    "assignments": [
                        {
                            "title": "Final Exam",
                            "due_date": "",
                            "points": "",
                            "description": "Due date and details to be announced in class and on Canvas."
                        }
                    ]
                },
                {
                    "name": "World War Two Essay",
                    "weight": "10",
                    "description": (
                        "This essay requires students to use primary sources (marked with a [P] in the schedule) to develop their own arguments about US involvement in World War Two. "
                        "This approximately 1200-word essay is due on Friday, March 14, at 9:30am EDT.\n\n"
                        "A. World War II Paper Instructions: This essay should include three parts:\n"
                        "• An introductory paragraph that clearly states your argument about whether the World War Two era was a high or low point of US commitment to freedom and democracy.\n"
                        "• Several paragraphs that substantiate your argument with specific, cited evidence from Chapter 5 of Gerstle’s book, American Crucible, and at least three primary sources.\n"
                        "• A 1-2 paragraph conclusion that summarizes your main argument and considers your argument’s implication for your broader understanding of the US government, democracy, warfare, personal liberty, and/or other relevant historical themes.\n\n"
                        "B. World War II Paper Rubric: This essay will be evaluated on a 0-100 scale using these criteria:\n"
                        "(20/100 points) Makes a clear argument in the introduction and substantiates that argument successfully throughout the paper.\n"
                        "(40/100 points) Paper demonstrates thoughtful engagement with several historical examples and draws from the Gerstle chapter and at least three primary sources.\n"
                        "(15/100 points) Conclusion summarizes the main argument and seriously considers that argument’s implication for your understanding of the federal government, democracy, warfare, individual liberty, and/or other relevant historical themes.\n"
                        "(15/100 points) Paper properly cites relevant primary and/or secondary literature using MLA, APA, or Chicago citations.\n"
                        "(10/100 points) Prose is clear, smoothly written, free of grammatical errors, and meets the word count base of 900 words (excluding title, citations, pagination, and bibliography)."
                    ),
                    "assignments": [
                        {
                            "title": "World War Two Essay",
                            "due_date": "Friday, March 14, 9:30am EDT",
                            "points": "",
                            "description": "See instructions and rubric above."
                        }
                    ]
                },
                {
                    "name": "Mini-Book Review",
                    "weight": "15",
                    "description": (
                        "Students will use primary sources and skills learned in the course to review the third and fourth chapters of Nancy MacLean’s, Freedom Is Not Enough (2006). "
                        "This approximately 1200-word paper is due on Friday, April 4, at 9:30am EDT.\n\n"
                        "Mini-Book Review Instructions: Your mini-book review should include four parts:\n"
                        "• One introductory paragraph that quickly (2-4 sentences) summarizes the chapters from Freedom Is Not Enough before clearly stating your argument about the book’s strengths and weaknesses in the last 1-2 sentences of that paragraph.\n"
                        "• A second paragraph that provides a brief but thorough synopsis of the MacLean chapter’s structure, main argument(s), and key examples or pieces of evidence.\n"
                        "• At least two body paragraphs that dissect the strengths and weaknesses of the book chapters by putting them into conversation with the at least three primary sources, at least two of which must be from the UF archives.\n"
                        "• A few sentences at the end of the review that provide your overall appraisal of the book section and your suggestion whether or not others should read it.\n\n"
                        "Mini-Book Review Grading Rubric: This essay will be graded on a 0-100 scale using the following criteria:\n"
                        "(30/100 points) First and second paragraphs accurately summarizes the book chapter and makes a clear argument about its strengths and/or weaknesses.\n"
                        "(40/100 points) Paper demonstrates thoughtful engagement with several historical examples and draws from at least three primary sources, two of which are from the UF archives.\n"
                        "(15/100 points) Concluding sentences provide a compelling, summative appraisal of the chapter and clear recommendation about whether or not others should read it.\n"
                        "(15/100 points) Paper properly cites relevant primary and/or secondary literature using MLA, APA, or Chicago citations.\n"
                        "(10 points) Prose is clear, smoothly written, free of grammatical errors, and meets the word count base of 1000 words (excluding title, citations, pagination, and bibliography)."
                    ),
                    "assignments": [
                        {
                            "title": "Mini-Book Review",
                            "due_date": "Friday, April 4, 9:30am EDT",
                            "points": "",
                            "description": "See instructions and rubric above."
                        }
                    ]
                }
            ]
            amh2020_template.schedule = []  # Leave schedule empty for user to fill in
            # Set policies
            amh2020_template.optional_policies = {
                "in_class_recording": True,
                "conflict_resolution": True,
                "campus_resources": True,
                "academic_resources": True,
                "late_work": True
            }
            # Add to templates list
            return [test_template, amh2020_template]
        except Exception as e:
            print(f"Error loading default templates: {e}")
            return []

    def create_action_buttons(self):
        """Create action buttons at the bottom of the interface"""
        action_frame = ttk.Frame(self.main_container)
        action_frame.pack(fill=tk.X, padx=5, pady=10)

        # Generate syllabus button (Word)
        gen_syllabus_btn = ttk.Button(
            action_frame, text="Generate Syllabus (Word)", style="Action.TButton",
            command=lambda: self.generate_syllabus(export_format="docx")
        )
        gen_syllabus_btn.pack(side=tk.RIGHT, padx=5)

        # Generate PDF button
        gen_pdf_btn = ttk.Button(
            action_frame, text="Generate Syllabus (PDF)", style="Action.TButton",
            command=lambda: self.generate_syllabus(export_format="pdf")
        )
        gen_pdf_btn.pack(side=tk.RIGHT, padx=5)

    def create_scrollable_frame(self, parent):
        """Create a scrollable frame within the given parent"""
        canvas = tk.Canvas(parent)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        # Add mousewheel scrolling
        self.add_mousewheel_scrolling(canvas, canvas)
        self.add_mousewheel_scrolling(scrollable_frame, canvas)
        
        return scrollable_frame

    def create_course_info_tab(self):
        """Create the course information tab"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Course Information")
        frame = self.create_scrollable_frame(tab)
        self.add_mousewheel_scrolling(tab, tab)

        # Course info fields
        fields = [
            ("Course Number:", "course_num", 15),
            ("Course Title:", "course_title", 40),
            ("Term:", "term", 20),
            ("Credits:", "credits", 5),
            ("Prerequisites:", "prerequisites", 40),
            ("Meeting Days/Times:", "meeting_times", 40),
            ("Location:", "location", 30)
        ]
        current_row = 0
        for label, field_name, width in fields:
            ttk.Label(frame, text=label).grid(row=current_row, column=0, sticky="e", padx=5, pady=5)
            entry = ttk.Entry(frame, width=width)
            entry.grid(row=current_row, column=1, sticky="w", padx=5, pady=5)
            setattr(self, f"entry_{field_name}", entry)
            current_row += 1

        # General Education (fixed to Social and Behavioral Sciences (S))
        ttk.Label(frame, text="General Education:").grid(row=current_row, column=0, sticky="e", padx=5, pady=5)
        gen_ed_frame = ttk.Frame(frame)
        gen_ed_frame.grid(row=current_row, column=1, sticky="w", padx=5, pady=5)
        ttk.Label(gen_ed_frame, text="Social and Behavioral Sciences (S)", font=("Arial", 10)).pack(side=tk.LEFT, padx=5)
        # Add checkbox for showing Gen Ed section in exports/preview
        gen_ed_check = ttk.Checkbutton(
            gen_ed_frame,
            text="Show General Education Designation in syllabus",
            variable=self.show_gen_ed,
            command=self.update_document_preview
        )
        gen_ed_check.pack(side=tk.LEFT, padx=10)
        current_row += 1

        # Course description
        ttk.Label(frame, text="Course Description:").grid(row=current_row, column=0, sticky="ne", padx=5, pady=5)
        self.txt_description = scrolledtext.ScrolledText(frame, width=60, height=6, wrap=tk.WORD)
        self.txt_description.grid(row=current_row, column=1, sticky="w", padx=5, pady=5)
        self.add_mousewheel_scrolling(self.txt_description)
        current_row += 1

        # Course Objectives as numbered list input
        ttk.Label(frame, text="Course Objectives:").grid(row=current_row, column=0, sticky="nw", padx=5, pady=5)
        self.objective_entries_frame = ttk.Frame(frame)
        self.objective_entries_frame.grid(row=current_row, column=1, sticky="w", padx=5, pady=5)
        self.objective_entries = []  # Start with an empty list so teachers can enter objectives
        add_objective_btn = ttk.Button(frame, text="Add Objective", command=self.add_objective_entry, style="Action.TButton")
        add_objective_btn.grid(row=current_row+1, column=1, sticky="w", padx=5, pady=5)
        current_row += 2

    def create_learning_objectives_tab(self):
        """Create a tab for customizable learning objectives table"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Learning Objectives & SLOs")
        self.add_mousewheel_scrolling(tab, tab)

        # Create vertical paned window to split content
        main_paned = ttk.PanedWindow(tab, orient=tk.VERTICAL)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Top section - Student Learning Outcomes
        top_frame = ttk.LabelFrame(main_paned, text="Student Learning Outcomes")
        main_paned.add(top_frame, weight=1)
        
        ttk.Label(top_frame, text="A student who successfully completes this course will:", 
                 style="Italic.TLabel").pack(anchor="w", padx=10, pady=(5,10))
        
        # Container for SLO entries
        self.outcome_entries_frame = ttk.Frame(top_frame)
        self.outcome_entries_frame.pack(fill=tk.X, expand=True, padx=10)
        self.outcome_entries = []
        # Add default SLO entries with empty content
        default_outcomes = [
            "",  # Empty strings for blank entries
            "",
            "",
            ""
        ]
        # Create initial 4 empty outcome entries
        for outcome in default_outcomes:
            self.add_outcome_entry(outcome)
        
        # Add button for new outcomes
        ttk.Button(top_frame, text="Add Learning Outcome", 
                   command=self.add_outcome_entry,
                   style="Action.TButton").pack(pady=10, padx=10)
        
        # Bottom section - Learning Objectives Table
        bottom_paned = ttk.PanedWindow(main_paned, orient=tk.HORIZONTAL)
        main_paned.add(bottom_paned, weight=2)
        
        # Left side - editing
        edit_frame = ttk.LabelFrame(bottom_paned, text="Learning Objectives Table")
        bottom_paned.add(edit_frame, weight=1)
        
        ttk.Label(edit_frame, text="Fill in the course-specific information for each category:", 
                 style="Italic.TLabel").pack(anchor="w", padx=10, pady=(5,10))
        
        # Scrollable container for learning objectives
        canvas = tk.Canvas(edit_frame)
        scrollbar = ttk.Scrollbar(edit_frame, orient="vertical", command=canvas.yview)
        
        self.lo_entries_frame = ttk.Frame(canvas)
        self.lo_entries_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=self.lo_entries_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack canvas and scrollbar
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Add default categories with blank content
        default_categories = {
            "Content": {
                "slo": "",
                "assignments": self.get_outcomes_range(),  # Dynamic based on number of outcomes
                "course_specific": ""
            },
            "Critical Thinking": {
                "slo": "",
                "assignments": self.get_outcomes_range(),
                "course_specific": ""
            },
            "Communication": {
                "slo": "",
                "assignments": self.get_outcomes_range(),
                "course_specific": ""
            }
        }
        for category, content in default_categories.items():
            self.add_learning_objective_row(category, content["slo"], content["assignments"])
        
        ttk.Button(edit_frame, text="Add New Category", 
                   command=lambda: self.add_learning_objective_row(),
                   style="Action.TButton").pack(pady=10)
        
        # Right side - preview
        preview_frame = ttk.LabelFrame(bottom_paned, text="Preview")
        bottom_paned.add(preview_frame, weight=1)
        
        self.preview_frame = ttk.Frame(preview_frame)
        self.preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Initial preview
        self.update_lo_preview()

    def create_instructor_info_tab(self):
        """Create the instructor information tab with proper layout"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Instructor Information")
        self.add_mousewheel_scrolling(tab, tab)

        # Create a canvas with scrollbar
        canvas = tk.Canvas(tab)
        scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        
        # Create main content frame
        content_frame = ttk.Frame(canvas)
        content_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        
        # Configure canvas
        canvas.create_window((0, 0), window=content_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack canvas and scrollbar
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Instructor Information
        instructor_frame = ttk.LabelFrame(content_frame, text="Instructor Information")
        instructor_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # Grid for instructor info with consistent field names
        fields = [
            ("Name:", "name", 40),
            ("Office:", "office", 40),
            ("Phone:", "phone", 20),
            ("Email:", "email", 40),
            ("Office Hours:", "office_hours", 60)
        ]
        for i, (label, field, width) in enumerate(fields):
            ttk.Label(instructor_frame, text=label).grid(row=i, column=0, sticky="e", padx=5, pady=5)
            entry = ttk.Entry(instructor_frame, width=width)
            entry.grid(row=i, column=1, sticky="w", padx=5, pady=5)
            setattr(self, f"entry_instr_{field}", entry)
        
        # Teaching Assistants
        self.ta_frame = ttk.LabelFrame(content_frame, text="Teaching Assistants")
        self.ta_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.ta_container = ttk.Frame(self.ta_frame)
        self.ta_container.pack(fill=tk.X, padx=5, pady=5)
        self.ta_entries = []
        
        ttk.Button(self.ta_frame, text="Add Teaching Assistant", 
                  command=self.add_ta,
                  style="Action.TButton").pack(pady=5)
        
    def create_schedule_tab(self):
        """Create the course schedule tab with proper layout"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Calendar/Course Schedule") # Renamed tab
        self.add_mousewheel_scrolling(tab, tab)

        # Main container
        main_frame = ttk.Frame(tab, padding="5")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Top controls frame
        controls_frame = ttk.Frame(main_frame)
        controls_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Left side buttons
        buttons_frame = ttk.Frame(controls_frame)
        buttons_frame.pack(side=tk.LEFT)
        
        ttk.Button(buttons_frame, text="Add Entry", 
                  command=self.add_schedule_entry,
                  style="Action.TButton").pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(buttons_frame, text="Import Schedule",
                  command=self.import_schedule,
                  style="Action.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Export Schedule",
                  command=self.export_schedule,
                  style="Action.TButton").pack(side=tk.LEFT, padx=5)
        
        # Right side helper text
        ttk.Label(controls_frame, text="Use [P] to mark primary sources", 
                  style="Italic.TLabel").pack(side=tk.RIGHT)
        
        # Fixed header frame for column headings
        header_frame = ttk.Frame(main_frame, style="Preview.TFrame")
        header_frame.pack(fill=tk.X, pady=(0, 5))  # Use pack instead of grid
        
        # Define column headings
        ttk.Label(header_frame, text="Date", style="Heading.TLabel").pack(side=tk.LEFT, padx=(5, 10))
        ttk.Label(header_frame, text="Topic", style="Heading.TLabel").pack(side=tk.LEFT, padx=150)
        ttk.Label(header_frame, text="Readings/Preparation", style="Heading.TLabel").pack(side=tk.LEFT, padx=5)
        ttk.Label(header_frame, text="Work Due", style="Heading.TLabel").pack(side=tk.LEFT, padx=175)
        
        # Scrollable frame for schedule entries
        canvas = tk.Canvas(main_frame, bg='#f5f5f5')  # Light gray background
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        
        # Create a frame inside the canvas for entries
        self.entries_frame = ttk.Frame(canvas, style="Template.TFrame")
        self.entries_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        # Configure the canvas
        canvas.create_window((0, 0), window=self.entries_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Pack the canvas and scrollbar
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=5)

        # Enable mousewheel scrolling
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        self.schedule_entries = []
        
    def create_assignments_tab(self):
        """Create the assignments and grading tab with proper layout"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="★ Assignments & Grading [Required]")
        self.add_mousewheel_scrolling(tab, tab)

        # Create a canvas with scrollbar for the main content
        canvas = tk.Canvas(tab)
        scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        
        # Create main content frame
        content_frame = ttk.Frame(canvas)
        content_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        
        # Configure canvas
        canvas.create_window((0, 0), window=content_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack canvas and scrollbar
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Required Materials Section
        materials_frame = ttk.LabelFrame(content_frame, text="Required Materials")
        materials_frame.pack(fill=tk.X, padx=5, pady=5)

        # --- Formatting Toolbar ---
        toolbar = ttk.Frame(materials_frame)
        toolbar.pack(fill=tk.X, padx=5, pady=(5, 0))

        def insert_markup(markup_type):
            text_widget = self.materials_text
            try:
                sel_start = text_widget.index(tk.SEL_FIRST)
                sel_end = text_widget.index(tk.SEL_LAST)
                selected = text_widget.get(sel_start, sel_end)
            except tk.TclError:
                sel_start = sel_end = None
                selected = ""
            if markup_type == "bold":
                tag = "**"
                new_text = f"{tag}{selected or 'bold text'}{tag}"
            elif markup_type == "italic":
                tag = "*"
                new_text = f"{tag}{selected or 'italic text'}{tag}"
            elif markup_type == "link":
                new_text = f"[{selected or 'link text'}](http://example.com)"
            else:
                return
            if sel_start and sel_end:
                text_widget.delete(sel_start, sel_end)
                text_widget.insert(sel_start, new_text)
            else:
                text_widget.insert(tk.INSERT, new_text)

        bold_btn = ttk.Button(toolbar, text="B", width=2, command=lambda: insert_markup("bold"))
        bold_btn.pack(side=tk.LEFT, padx=(0, 2))
        italic_btn = ttk.Button(toolbar, text="I", width=2, command=lambda: insert_markup("italic"))
        italic_btn.pack(side=tk.LEFT, padx=(0, 2))
        link_btn = ttk.Button(toolbar, text="🔗", width=2, command=lambda: insert_markup("link"))
        link_btn.pack(side=tk.LEFT, padx=(0, 2))
        ttk.Label(toolbar, text="(Use toolbar or Markdown: *italic* **bold** [text](url))", font=("Arial", 8, "italic")).pack(side=tk.LEFT, padx=8)

        self.materials_text = scrolledtext.ScrolledText(materials_frame, width=60, height=4, wrap=tk.WORD)
        self.materials_text.pack(fill=tk.X, padx=5, pady=5)
        
        fee_frame = ttk.Frame(materials_frame)
        fee_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(fee_frame, text="Materials Fee: $").pack(side=tk.LEFT)
        self.fee_entry = ttk.Entry(fee_frame, width=10)
        self.fee_entry.pack(side=tk.LEFT)
        
        # Add this to the relevant section of your UI (e.g., in the Required Materials section)
        help_button = ttk.Button(materials_frame, text="Formatting Help", command=self.show_formatting_help)
        help_button.pack(side=tk.RIGHT, padx=5, pady=5)

        # Grading Components Section
        components_frame = ttk.LabelFrame(content_frame, text="Graded Components")
        components_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.categories_frame = ttk.Frame(components_frame)
        self.categories_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.category_frames = []
        
        ttk.Button(components_frame, text="Add Category", 
                  command=self.add_category,
                  style="Action.TButton").pack(padx=5, pady=5)
        
        # Grading Scale Section
        scale_frame = ttk.LabelFrame(content_frame, text="Grading Scale")
        scale_frame.pack(fill=tk.X, padx=5, pady=5)
        
        grades_frame = ttk.Frame(scale_frame)
        grades_frame.pack(padx=5, pady=5)
        
        grades = [
            ("A", "93-100"), ("A-", "90-92"),
            ("B+", "87-89"), ("B", "83-86"),
            ("B-", "80-82"), ("C+", "77-79"),
            ("C", "73-76"), ("C-", "70-72"),
            ("D+", "67-69"), ("D", "63-66"),
            ("D-", "60-62"), ("E", "0-59")
        ]
        
        for i, (letter, number) in enumerate(grades):
            row = i // 4
            col = i % 4
            grade_frame = ttk.Frame(grades_frame)
            grade_frame.grid(row=row, column=col, padx=5, pady=2)
            ttk.Label(grade_frame, text=f"{letter}:").pack(side=tk.LEFT)
            entry = ttk.Entry(grade_frame, width=8)
            entry.insert(0, number)
            entry.pack(side=tk.LEFT, padx=2)
        
    def create_policies_tab(self):
        """Create the policies tab"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Policies")
        self.add_mousewheel_scrolling(tab, tab)
        
        frame = self.create_scrollable_frame(tab)
        
        # University Assessment Policies
        self.assessment_policy_var = tk.BooleanVar(value=True) # Default to True
        assessment_check = ttk.Checkbutton(
            frame,
            text="University Assessment Policies",
            variable=self.assessment_policy_var,
            command=self.update_document_preview
        )
        assessment_check.pack(anchor="w", padx=20, pady=2)
        
        # Extensions & Make-Up Exams
        self.extensions_policy_var = tk.BooleanVar(value=True) # Default to True
        extensions_check = ttk.Checkbutton(
            frame,
            text="Extensions & Make-Up Exams",
            variable=self.extensions_policy_var,
            command=self.update_document_preview
        )
        extensions_check.pack(anchor="w", padx=20, pady=2)
        
        # Canvas Policy section
        canvas_frame = ttk.LabelFrame(frame, text="Canvas Policy")
        canvas_frame.pack(fill=tk.X, padx=20, pady=(10, 5))
        
        self.canvas_policy_text = scrolledtext.ScrolledText(canvas_frame, width=60, height=4, wrap=tk.WORD)
        self.canvas_policy_text.pack(fill=tk.X, padx=5, pady=5)
        self.canvas_policy_text.insert("1.0", "-Replace with your Canvas Policies-")
        self.add_mousewheel_scrolling(self.canvas_policy_text)
        
        ttk.Label(canvas_frame, text="This will appear under the Extra Credit policy in the syllabus.", 
                 style="Italic.TLabel").pack(anchor="w", padx=5, pady=(0, 5))
        
        # Technology in the Classroom Policy
        tech_frame = ttk.LabelFrame(frame, text="Technology in the Classroom Policy")
        tech_frame.pack(fill=tk.X, padx=20, pady=(10, 5))
        
        self.technology_policy_text = scrolledtext.ScrolledText(tech_frame, width=60, height=4, wrap=tk.WORD)
        self.technology_policy_text.pack(fill=tk.X, padx=5, pady=5)
        self.technology_policy_text.insert("1.0", "-Replace with your Technology Policies-")
        self.add_mousewheel_scrolling(self.technology_policy_text)
        
        ttk.Label(tech_frame, text="This will appear after the Canvas policy in the syllabus.", 
                 style="Italic.TLabel").pack(anchor="w", padx=5, pady=(0, 5))
        
        # Class Communication Policy
        comm_frame = ttk.LabelFrame(frame, text="Class Communication Policy")
        comm_frame.pack(fill=tk.X, padx=20, pady=(10, 5))
        
        self.communication_policy_text = scrolledtext.ScrolledText(comm_frame, width=60, height=4, wrap=tk.WORD)
        self.communication_policy_text.pack(fill=tk.X, padx=5, pady=5)
        self.communication_policy_text.insert("1.0", "-Replace with your policy of communication with students-")
        self.add_mousewheel_scrolling(self.communication_policy_text)
        
        ttk.Label(comm_frame, text="This will appear after the Technology policy in the syllabus.", 
                 style="Italic.TLabel").pack(anchor="w", padx=5, pady=(0, 5))
        
        # Optional policies (moved below comm)
        support_check = ttk.Checkbutton(
            frame,
            text="Assignment Support Outside the Classroom",
            variable=self.outside_support_var,
            command=self.update_document_preview
        )
        support_check.pack(anchor="w", padx=20, pady=2)
        support_frame = ttk.LabelFrame(frame, text="Assignment Support Details")
        support_frame.pack(fill=tk.X, padx=20, pady=(10, 5))
        self.support_text = scrolledtext.ScrolledText(support_frame, width=60, height=4, wrap=tk.WORD)
        self.support_text.pack(fill=tk.X, padx=5, pady=5)
        self.support_text.insert("1.0", "-Write your policy for assisting students outside of the classroom here-")
        
        # Fixed policies (non-editable) -> Now optional checkboxes
        checkbox_in_class_recording = ttk.Checkbutton(
            frame,
            text="In-class recording policy",
            variable=self.in_class_recording_var,
            command=self.update_document_preview
        )
        checkbox_in_class_recording.pack(anchor="w", padx=5, pady=2)

        # Optional: Procedure for Conflict Resolution
        checkbox_conflict_resolution = ttk.Checkbutton(
            frame,
            text="Procedure for Conflict Resolution",
            variable=self.conflict_resolution_var,
            command=self.update_document_preview
        )
        checkbox_conflict_resolution.pack(anchor="w", padx=5, pady=2)

        # Optional: Campus Resources
        checkbox_campus_resources = ttk.Checkbutton(
            frame,
            text="Campus Resources",
            variable=self.campus_resources_var,
            command=self.update_document_preview
        )
        checkbox_campus_resources.pack(anchor="w", padx=5, pady=2)

        # Optional: Academic Resources
        checkbox_academic_resources = ttk.Checkbutton(
            frame,
            text="Academic Resources",
            variable=self.academic_resources_var,
            command=self.update_document_preview
        )
        checkbox_academic_resources.pack(anchor="w", padx=5, pady=2)

        # --- Move Late Submissions Policy Section here ---
        late_frame = ttk.LabelFrame(frame, text="Late Submissions Policy")
        late_frame.pack(fill=tk.X, padx=20, pady=5)
        ttk.Label(late_frame, text="Late Submissions Policy:").pack(side=tk.LEFT, padx=(0, 5))
        
        self.late_policy_var = tk.StringVar()
        self.late_policies = {
            "Standard (10% per day)": "Unless an extension is granted, assignments will incur a 10-point penalty for every day they are late.",
            "No late work": "No late work will be accepted without prior approval.",
            "48-hour grace": "Students have a 48-hour grace period for submissions, after which no late work will be accepted.",
            "Custom": ""
        }
        self.late_combo = ttk.Combobox(late_frame, 
                                      textvariable=self.late_policy_var,
                                      values=list(self.late_policies.keys()), 
                                      width=30)
        self.late_combo.pack(side=tk.LEFT)
        self.late_combo.set("Select a policy...")
        
        self.late_policy_text = scrolledtext.ScrolledText(late_frame, width=60, height=3, wrap=tk.WORD)
        self.late_policy_text.pack(fill=tk.X, padx=5, pady=(0, 5))
        
        def update_late_policy(*args):
            selected = self.late_policy_var.get()
            if selected in self.late_policies:
                # Enable text widget
                self.late_policy_text.config(state='normal')
                self.late_policy_text.delete('1.0', tk.END)
                self.late_policy_text.insert('1.0', self.late_policies[selected])
                if selected != "Custom":
                    self.late_policy_text.config(state='disabled')
                # Update optional policy state just by setting it:
                if hasattr(self, 'optional_policies'):
                    # Check if the key exists before setting, although it should now
                    if "late_work" in self.optional_policies:
                        if selected == "No late work":
                            self.optional_policies["late_work"].set(False)
                        else:
                            self.optional_policies["late_work"].set(True)
        
        self.late_policy_var.trace_add('write', update_late_policy)
        self.late_combo.bind('<<ComboboxSelected>>', lambda e: update_late_policy())

        # --- Move Extra Credit Policy Section here ---
        extra_frame = ttk.LabelFrame(frame, text="Extra Credit Policy")
        extra_frame.pack(fill=tk.X, padx=20, pady=5)
        ttk.Label(extra_frame, text="Extra Credit Policy:").pack(side=tk.LEFT, padx=(0, 5))
        
        self.extra_credit_var = tk.StringVar()
        self.extra_credit_policies = {
            "Standard": "Extra credit opportunities may be announced during the semester. Points will be added to your mid-term exam grade.",
            "No extra credit": "No extra credit will be offered in this course.",
            "Optional assignments": "Students may complete optional assignments for extra credit, worth up to 3% of the final grade.",
            "Custom": ""
        }
        self.extra_combo = ttk.Combobox(extra_frame, textvariable=self.extra_credit_var,
                                        values=list(self.extra_credit_policies.keys()), width=30)
        self.extra_combo.pack(side=tk.LEFT)
        self.extra_combo.set("Select a policy...")
        
        self.extra_credit_text = scrolledtext.ScrolledText(extra_frame, width=60, height=3, wrap=tk.WORD)
        self.extra_credit_text.pack(fill=tk.X, padx=5, pady=(0, 5))
        
        def update_extra_credit(*args):
            selected = self.extra_credit_var.get()
            if selected in self.extra_credit_policies:
                self.extra_credit_text.config(state='normal')
                self.extra_credit_text.delete('1.0', tk.END)
                self.extra_credit_text.insert('1.0', self.extra_credit_policies[selected])
                if selected != "Custom":
                    self.extra_credit_text.config(state='disabled')
        
        self.extra_credit_var.trace_add('write', update_extra_credit)
        self.extra_combo.bind('<<ComboboxSelected>>', lambda e: update_extra_credit())

    def create_action_buttons(self):
        """Create action buttons at the bottom of the interface"""
        action_frame = ttk.Frame(self.main_container)
        action_frame.pack(fill=tk.X, padx=5, pady=10)


        # Generate syllabus button (Word)
        gen_syllabus_btn = ttk.Button(
            action_frame, text="Generate Syllabus (Word)", style="Action.TButton",
            command=lambda: self.generate_syllabus(export_format="docx")
        )
        gen_syllabus_btn.pack(side=tk.RIGHT, padx=5)
        # Generate PDF button
        gen_pdf_btn = ttk.Button(
            action_frame, text="Generate Syllabus (PDF)", style="Action.TButton",
            command=lambda: self.generate_syllabus(export_format="pdf")
        )
        gen_pdf_btn.pack(side=tk.RIGHT, padx=5)
        
    def add_ta(self):
        """Add a new TA entry"""
        frame = ttk.Frame(self.ta_container)
        frame.pack(anchor="w", pady=2)
        
        entries = []
        for label, width in [("Name:", 25), ("Email:", 30), ("Office Hours:", 45)]:
            ttk.Label(frame, text=label).pack(side=tk.LEFT)
            entry = ttk.Entry(frame, width=width)
            entry.pack(side=tk.LEFT, padx=5)
            entries.append(entry)
            
        self.ta_entries.append(entries)
        
        def remove_ta():
            frame.destroy()
            self.ta_entries.remove(entries)
            
        remove_btn = ttk.Button(frame, text="X", 
                              command=remove_ta,
                              style="Delete.TButton")
        remove_btn.pack(side=tk.LEFT, padx=2)
        
    def add_schedule_entry(self, date="", topic="", readings="", work_due=""):
        """Add a new schedule entry with enhanced styling"""
        row = len(self.schedule_entries)
        
        # Date entry
        date_entry = ttk.Entry(self.entries_frame, width=15)
        date_entry.grid(row=row, column=0, sticky="w", padx=(5, 10), pady=2)
        date_entry.insert(0, date)  # Populate with provided data
        
        # Topic entry
        topic_entry = ttk.Entry(self.entries_frame, width=40)
        topic_entry.grid(row=row, column=1, sticky="ew", padx=5, pady=2)
        topic_entry.insert(0, topic)  # Populate with provided data
        
        # Readings frame
        readings_frame = ttk.Frame(self.entries_frame)
        readings_frame.grid(row=row, column=2, sticky="ew", padx=5, pady=2)
        readings_frame.grid_columnconfigure(0, weight=1)
        
        # Readings text area
        readings_text = scrolledtext.ScrolledText(readings_frame, width=50, height=3, wrap=tk.WORD)
        readings_text.grid(row=0, column=0, sticky="ew")
        readings_text.insert("1.0", readings)  # Populate with provided data
        self.add_mousewheel_scrolling(readings_text)
        
        # Buttons frame
        buttons_frame = ttk.Frame(readings_frame)
        buttons_frame.grid(row=0, column=1, sticky="ns", padx=(5, 0))
        
        def insert_p_marker():
            readings_text.insert(tk.INSERT, "[P] ")
        
        def count_words():
            text = readings_text.get("1.0", tk.END).strip()
            word_count = len(text.split())
            readings_text.insert(tk.END, f" [{word_count} words]")
        
        ttk.Button(buttons_frame, text="[P]", command=insert_p_marker, style="Small.TButton").pack(side=tk.TOP, pady=(0, 2))
        ttk.Button(buttons_frame, text="#", command=count_words, style="Small.TButton").pack(side=tk.TOP)
        
        # Work Due entry
        work_due_entry = ttk.Entry(self.entries_frame, width=20)
        work_due_entry.grid(row=row, column=3, sticky="w", padx=5, pady=2)
        work_due_entry.insert(0, work_due)  # Populate with provided data
        
        def remove_entry():
            date_entry.destroy()
            topic_entry.destroy()
            readings_frame.destroy()
            work_due_entry.destroy()
            delete_btn.destroy()
            self.schedule_entries.remove(entry_dict)
            self.repack_schedule_entries()
        
        # Delete button
        delete_btn = ttk.Button(self.entries_frame, text="X", command=remove_entry, style="Delete.TButton")
        delete_btn.grid(row=row, column=4, padx=(0, 5), pady=2)
        
        entry_dict = {
            "date": date_entry,
            "topic": topic_entry,
            "readings": readings_text,
            "work_due": work_due_entry,
            "delete_btn": delete_btn,
            "row": row
        }
        
        self.schedule_entries.append(entry_dict)
        return entry_dict

    def repack_schedule_entries(self):
        """Repack all schedule entries after a deletion"""
        for i, entry in enumerate(self.schedule_entries):
            entry["row"] = i  # Update row number
            entry["date"].grid(row=i, column=0, sticky="w", padx=(5, 10), pady=2)
            entry["topic"].grid(row=i, column=1, sticky="ew", padx=5, pady=2)
            entry["readings"].master.grid(row=i, column=2, sticky="ew", padx=5, pady=2)
            entry["work_due"].grid(row=i, column=3, sticky="w", padx=5, pady=2)
            entry["delete_btn"].grid(row=i, column=4, padx=(0, 5), pady=2)
        
    def on_template_selected(self, event):
        """Handle template selection. If 'Clear Template' is selected, reset all fields."""
        try:
            selected = self.template_combo.get().strip()
            if selected == "Clear Template":
                # Reset all fields to default state
                self.clear_all_entries()
                # Optionally, reinitialize default empty outcomes in the Learning Objectives tab:
                default_outcomes = ["", "", "", ""]
                for outcome in default_outcomes:
                    self.add_outcome_entry(outcome)
                # Also clear any objectives entries if desired
                for obj in self.objective_entries:
                    obj["entry"].delete(0, tk.END)
                # Reset the combo box itself
                self.template_var.set("")
                # Update previews
                self.update_lo_preview()
                self.update_document_preview()
                return

            # Otherwise, load the selected template as before
            template = None
            for t in self.templates:
                if t.title == selected or f"{t.course_code}: {t.title}" == selected:
                    template = t
                    break

            if template:
                self.clear_all_entries()
                self.load_template_content(template)
                self.update_lo_preview()
                self.update_document_preview()
            else:
                print(f"Template not found: {selected}")

        except Exception as e:
            print(f"Error in template selection: {e}")
            import traceback
            traceback.print_exc()
    
    def load_template_content(self, template):
        """Load template content into form fields"""
        try:
            # Clear existing content first
            self.clear_all_entries()
            
            # Course Info
            if hasattr(template, 'course_code'):
                self.entry_course_num.insert(0, template.course_code)
            if hasattr(template, 'title'):
                self.entry_course_title.insert(0, template.title)
            if hasattr(template, 'prerequisites'):
                self.entry_prerequisites.insert(0, template.prerequisites)
            # Additional course info fields
            if hasattr(template, 'semester'):
                self.entry_term.insert(0, template.semester)
            if hasattr(template, 'credits'):
                self.entry_credits.insert(0, template.credits)
            if hasattr(template, 'class_days') and hasattr(template, 'class_times'):
                meeting_times = f"{template.class_days} {template.class_times}"
                self.entry_meeting_times.insert(0, meeting_times)
            if hasattr(template, 'classroom'):
                self.entry_location.insert(0, template.classroom)

            # Instructor Info
            if hasattr(template, 'instructor_name'):
                self.entry_instr_name.insert(0, template.instructor_name)
            if hasattr(template, 'instructor_office'):
                self.entry_instr_office.insert(0, template.instructor_office)
            if hasattr(template, 'instructor_phone'):
                self.entry_instr_phone.insert(0, template.instructor_phone)
            if hasattr(template, 'instructor_email'):
                self.entry_instr_email.insert(0, template.instructor_email)
            if hasattr(template, 'instructor_office_hours'):
                self.entry_instr_office_hours.insert(0, template.instructor_office_hours)
                
            # TAs
            if hasattr(template, 'tas') and template.tas:
                # Clear existing TAs first
                if hasattr(self, 'ta_entries'):
                    for ta_entries in self.ta_entries:
                        for entry in ta_entries:
                            entry.master.destroy()
                    self.ta_entries.clear()
                # Recreate TA container
                if hasattr(self, 'ta_container'):
                    self.ta_container.destroy()
                    self.ta_container = ttk.Frame(self.ta_frame)
                    self.ta_container.pack(fill=tk.X, padx=5, pady=5)
                
                # Add each TA from the template
                for ta_data in template.tas:
                    # Create new TA entry
                    self.add_ta()
                    # Get the latest TA entry (the one we just added)
                    if self.ta_entries:
                        latest_ta = self.ta_entries[-1]
                        latest_ta[0].insert(0, ta_data.get('name', ''))
                        latest_ta[1].insert(0, ta_data.get('email', ''))
                        latest_ta[2].insert(0, ta_data.get('office_hours', ''))
                    
            # Course Description & Objectives
            if hasattr(template, 'description'):
                self.txt_description.insert("1.0", template.description)
            if hasattr(template, 'objectives'):
                for obj in template.objectives:
                    self.add_objective_entry(obj)
                
            # Schedule
            if hasattr(template, 'schedule') and template.schedule:
                for entry in template.schedule:
                    self.add_schedule_entry(
                        date=entry.get('date', ''),
                        topic=entry.get('topic', ''),
                        readings=entry.get('readings', ''),
                        work_due=entry.get('work_due', '')
                    )
            
            # Grading Categories
            if hasattr(template, 'grading_categories') and template.grading_categories:
                for category_data in template.grading_categories:
                    category = self.add_category(
                        name=category_data.get('name', ''),
                        weight=category_data.get('weight', ''),
                        description=category_data.get('description', '')
                    )
                    # Add assignments for this category
                    if 'assignments' in category_data and category_data['assignments']:
                        for assignment in category_data['assignments']:
                            self.add_assignment_to_category(
                                category,
                                title=assignment.get('title', ''),
                                due_date=assignment.get('due_date', ''),
                                points=assignment.get('points', '')
                            )

            # Materials and Fee
            if hasattr(self, 'materials_text'):
                self.materials_text.delete('1.0', tk.END)
                self.materials_text.insert('1.0', "**Required** textbook *and* materials will be *listed* here.")
                
            if hasattr(self, 'fee_entry'):
                self.fee_entry.delete(0, tk.END)
                self.fee_entry.insert(0, "0.00")

            # Remove all existing categories before adding new ones
            for category in list(self.category_frames):
                if 'frame' in category:
                    category['frame'].destroy()
            self.category_frames.clear()

            # Handle categories based on template type
            if getattr(template, "title", "") == "TEST TEMPLATE":
                # Create 5 demo categories for test template
                for i in range(1, 6):
                    cat = self.add_category(
                        name=f"Component {i}",
                        weight="20",
                        description=f"This is the description for graded component {i}."
                    )
                    cat['name'].insert(0, f"Component {i}")
                    cat['weight'].insert(0, "20")
                    cat['description'].delete("1.0", tk.END)
                    cat['description'].insert("1.0", f"This is the description for graded component {i}.")
                    # Add 2 assignments per category for demonstration
                    for j in range(1, 3):
                        assignment = self.add_assignment_to_category(
                            cat,
                            title=f"Assignment {i}.{j}",
                        )
                        # Populate the assignment fields
                        assignment['title'].insert(0, f"Assignment {i}.{j}")
                        assignment['due date'].insert(0, f"2025-01-{i*2+j:02d}")
                        assignment['points'].insert(0, str(i * 10 + j))
                        assignment['description'].delete("1.0", tk.END)
                        assignment['description'].insert("1.0", f"Description for assignment {i}.{j}")
            # Handle regular template categories
            elif hasattr(template, 'grading_categories'):
                for category_data in template.grading_categories:
                    # Create new category
                    category = self.add_category(
                        name=category_data.get('name', ''),
                        weight=category_data.get('weight', ''),
                        description=category_data.get('description', '')
                    )
                    # Add assignments if they exist
                    if 'assignments' in category_data and category_data['assignments']:
                        for assignment in category_data['assignments']:
                            self.add_assignment_to_category(
                                category,
                                title=assignment.get('title', ''),
                                due_date=assignment.get('due_date', ''),
                                points=assignment.get('points', '')
                            )

            # Policies
            if hasattr(template, 'optional_policies') and template.optional_policies:
                for policy_name, enabled in template.optional_policies.items():
                    if policy_name in self.optional_policies:
                        self.optional_policies[policy_name].set(enabled)
            # Set dropdowns for late submissions and extra credit policy if present in template
            # Do this *after* the widgets are created and visible
            self.root.after(100, self._set_policy_dropdowns, template)

            # Load Student Learning Outcomes if provided in template
            if hasattr(template, 'outcomes') and template.outcomes:
                # Clear existing outcomes first
                for outcome_entry in list(self.outcome_entries):
                    outcome_entry["frame"].destroy()
                self.outcome_entries.clear()
                
                # Add each outcome from template
                for outcome in template.outcomes:
                    self.add_outcome_entry(outcome)
            
            # Load learning objectives with SLOs from template
            if hasattr(template, 'learning_objectives'):
                for category, content in template.learning_objectives.items():
                    if category in self.learning_objectives_entries:
                        entries = self.learning_objectives_entries[category]
                        if 'slo' in entries:
                            entries['slo'].config(state='normal')
                            entries['slo'].delete("1.0", tk.END)
                            entries['slo'].insert("1.0", content.get('slo', ''))
                            if category in ["Content", "Critical Thinking", "Communication"]:
                                entries['slo'].config(state='disabled')
                        
                        if 'assignments' in entries:
                            entries['assignments'].config(state='normal')
                            entries['assignments'].delete("1.0", tk.END)
                            entries['assignments'].insert("1.0", self.get_outcomes_range())
                            if category in ["Content", "Critical Thinking", "Communication"]:
                                entries['assignments'].config(state='disabled')

                        if 'course_specific' in entries:
                            entries['course_specific'].delete("1.0", tk.END)
                            entries['course_specific'].insert("1.0", content.get('course_specific', ''))

            # Update all references and previews
            self.update_outcomes_references()
            self.update_lo_preview()
            self.update_document_preview()
        except Exception as e:
            print(f"Error loading template content: {e}")
            import traceback
            traceback.print_exc()
    
    def _set_policy_dropdowns(self, template):
        """Set the late submissions and extra credit dropdowns after widgets are ready"""
        if hasattr(template, 'late_policy') and hasattr(self, 'late_policy_var'):
            self.late_policy_var.set(template.late_policy)
        if hasattr(template, 'extra_credit_policy') and hasattr(self, 'extra_credit_var'):
            self.extra_credit_var.set(template.extra_credit_policy)

    def clear_all_entries(self):
        """Clear all entries in the form"""
        # Clear course info (added self.entry_prerequisites)
        for entry in [
            self.entry_course_num, self.entry_course_title,
            self.entry_term, self.entry_credits,
            self.entry_prerequisites,
            self.entry_meeting_times, self.entry_location
        ]:
            entry.delete(0, tk.END)
        
        # Clear description
        self.txt_description.delete("1.0", tk.END)
        
        # Clear instructor info with correct field names
        for field in ["name", "office", "phone", "email", "office_hours"]:
            entry = getattr(self, f"entry_instr_{field}", None)
            if entry:
                entry.delete(0, tk.END)
        
        # Clear TAs
        if hasattr(self, 'ta_entries'):
            for ta_entries in self.ta_entries:
                for entry in ta_entries:
                    entry.master.destroy()
            self.ta_entries.clear()
        # Recreate TA container
        if hasattr(self, 'ta_container'):
            self.ta_container.destroy()
            self.ta_container = ttk.Frame(self.ta_frame)
            self.ta_container.pack(fill=tk.X, padx=5, pady=5)
        
        # Clear materials
        if hasattr(self, 'materials_text'):
            self.materials_text.delete('1.0', tk.END)
        if hasattr(self, 'fee_entry'):
            self.fee_entry.delete(0, tk.END)
        
        # Clear grading categories
        if hasattr(self, 'category_frames'):
            for category in self.category_frames:
                category["frame"].destroy()
            self.category_frames.clear()
        
        # Clear schedule entries
        if hasattr(self, 'schedule_entries'):
            for entry in self.schedule_entries:
                for widget in [entry["date"], entry["topic"], entry["readings"].master, entry["work_due"], entry["delete_btn"]]:
                    if widget.winfo_exists():
                        widget.destroy()
            self.schedule_entries.clear()

        # Clear outcome entries:
        if hasattr(self, 'outcome_entries'):
            for outcome_entry in list(self.outcome_entries):
                outcome_entry["frame"].destroy()
            self.outcome_entries = []
            # Optionally, re-add default outcomes if desired.

        # Clear course objectives (numbered list)
        if hasattr(self, 'objective_entries'):
            for obj in self.objective_entries:
                obj["frame"].destroy()    
            self.objective_entries = []

    def add_assignment_to_category(self, category, title="", due_date="", points=""):
        """Add a new assignment to a category"""
        assignment_frame = ttk.Frame(category["frame"])
        assignment_frame.pack(fill=tk.X, padx=5, pady=2)
        
        ttk.Label(assignment_frame, text="Title:").pack(side=tk.LEFT)
        title_entry = ttk.Entry(assignment_frame, width=30)
        title_entry.pack(side=tk.LEFT, padx=2)
        
        ttk.Label(assignment_frame, text="Due:").pack(side=tk.LEFT)
        due_entry = ttk.Entry(assignment_frame, width=15)
        due_entry.pack(side=tk.LEFT, padx=2)
        
        ttk.Label(assignment_frame, text="Points:").pack(side=tk.LEFT)
        points_entry = ttk.Entry(assignment_frame, width=5)
        points_entry.pack(side=tk.LEFT, padx=2)
        
        # Description for the assignment
        ttk.Label(assignment_frame, text="Description:").pack(side=tk.LEFT)
        description_text = scrolledtext.ScrolledText(assignment_frame, width=40, height=3, wrap=tk.WORD)
        description_text.pack(side=tk.LEFT, padx=2)
        
        def remove_assignment():
            assignment_frame.destroy()
            category["assignments"].remove(assignment_dict)
        
        remove_btn = ttk.Button(assignment_frame, text="X", 
                              command=remove_assignment,
                              style="Delete.TButton")
        remove_btn.pack(side=tk.LEFT, padx=2)
        
        assignment_dict = {
            "frame": assignment_frame,
            "title": title_entry,
            "due date": due_entry,
            "points": points_entry,
            "description": description_text  # Add description to the dictionary
        }
        
        if "assignments" not in category:
            category["assignments"] = []
        category["assignments"].append(assignment_dict)
        
        return assignment_dict

    def save_template(self):
        """Save current content as a template"""
        # Get content
        content = self.gather_content()
        
        # Get name
        name = simpledialog.askstring("Template Name", "Enter a name for this template:")
        if not name:
            return
        
        # Create template
        template = SyllabusTemplate(
            course_code=content["course_info"]["course_num"],
            title=content["course_info"]["course_title"],
            description=content["course_info"]["description"],
            objectives=content["course_info"]["objectives"].split('\n'),
            outcomes=[outcome.get("text") for outcome in content["outcomes"]]
        )
        # Add all content
        template.semester = content["course_info"]["term"]
        template.credits = content["course_info"]["credits"]
        template.prerequisites = content["course_info"].get("prerequisites", "")
        template.class_days = content["course_info"].get("meeting_days", "")
        template.class_times = content["course_info"].get("meeting_times", "")
        template.classroom = content["course_info"].get("location", "")
        # Instructor info
        template.instructor_name = content["instructor_info"].get("name", "")
        template.instructor_email = content["instructor_info"].get("email", "")
        template.instructor_office = content["instructor_info"].get("office", "")
        template.instructor_office_hours = content["instructor_info"].get("office_hours", "")
        template.instructor_phone = content["instructor_info"].get("phone", "")
        # Add TAs
        template.tas = content["tas"]
        # Add schedule
        template.schedule = content["schedule"]
        
        # Add grading categories
        template.grading_categories = content["grading_categories"]
        
        # Add policies
        template.optional_policies = content["optional_policies"]
        template.late_policy = content["late_policy"]
        template.extra_credit_policy = content["extra_credit_policy"]
        
        # Add learning objectives
        template.learning_objectives = content["learning_objectives"]
        
        # Store template in the list
        # Check if we already have a template with this name
        template_index = -1
        for i, existing_template in enumerate(self.templates):
            if hasattr(existing_template, 'course_code') and hasattr(existing_template, 'title'):
                existing_name = f"{existing_template.course_code}: {existing_template.title}"
                if existing_name == name:
                    template_index = i
                    break
        
        if template_index >= 0:
            # Update existing template
            self.templates[template_index] = template
        else:
            # Add as new template
            self.templates.append(template)
        
        # Update template names list
        self.template_names = [f"{t.course_code}: {t.title}" for t in self.templates]
        
        # Update template selector with new list
        self.template_combo['values'] = self.template_names
        
        messagebox.showinfo("Template Saved", f"Template '{name}' has been saved.")
        
        # Save templates to file for future use
        try:
            with open("syllabus_templates.pickle", "wb") as f:
                pickle.dump(self.templates, f)
        except Exception as e:
            messagebox.showwarning("Warning", f"Could not save templates to file: {str(e)}")
                
    def load_template(self):
        """Load templates from file"""
        try:
            with open("syllabus_templates.pickle", "rb") as f:
                loaded_templates = pickle.load(f)
                if loaded_templates:
                    # Make sure it's a list
                    if isinstance(loaded_templates, dict):
                        # Convert from old dictionary format to list format
                        self.templates = list(loaded_templates.values())
                    else:
                        self.templates = loaded_templates

                    # Update template names
                    self.template_names = [f"{t.course_code}: {t.title}" for t in self.templates]
                    # Update combobox
                    if hasattr(self, 'template_combo'):
                        self.template_combo['values'] = self.template_names
        except FileNotFoundError:
            # No templates file yet, just use what we have
            pass
        except Exception as e:
            messagebox.showwarning("Warning", f"Could not load templates: {str(e)}")
            
    def import_schedule(self):
        """Import schedule from CSV with enhanced formatting support"""
        file_path = filedialog.askopenfilename(
            filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")],
            title="Import Schedule"
        )
        if not file_path:
            return
        
        # Clear existing entries
        for entry in self.schedule_entries:
            entry["frame"].destroy()
        self.schedule_entries.clear()
        
        # Import based on file type
        if file_path.endswith('.csv'):
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    entry = self.add_schedule_entry()
                    entry["date"].insert(0, row.get("Date", ""))
                    entry["topic"].insert(0, row.get("Topic", ""))
                    entry["readings"].insert("1.0", row.get("Readings/Preparation", ""))
                    entry["work_due"].insert(0, row.get("Work Due", ""))
        else:
            # Handle Excel import
            import pandas as pd
            df = pd.read_excel(file_path)
            for _, row in df.iterrows():
                entry = self.add_schedule_entry()
                entry["date"].insert(0, str(row.get("Date", "")))
                entry["topic"].insert(0, str(row.get("Topic", "")))
                entry["readings"].insert("1.0", str(row.get("Readings/Preparation", "")))
                entry["work_due"].insert(0, str(row.get("Work Due", "")))
            
    def export_schedule(self):
        """Export schedule to CSV with formatting preserved"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")],
            title="Export Schedule"
        )
        if not file_path:
            return
        
        # Prepare data
        data = []
        for entry in self.schedule_entries:
            data.append({
                "Date": entry["date"].get(),
                "Topic": entry["topic"].get(),
                "Readings/Preparation": entry["readings"].get("1.0", tk.END).strip(),
                "Work Due": entry["work_due"].get()
            })
        
        # Export based on file type
        if file_path.endswith('.csv'):
            import csv
            with open(file_path, 'w', encoding='utf-8', newline='') as f:
                writer = csv.DictWriter(f, fieldnames=["Date", "Topic", "Readings/Preparation", "Work Due"])
                writer.writeheader()
                writer.writerows(data)
        else:
            # Handle Excel export
            import pandas as pd
            df = pd.DataFrame(data)
            df.to_excel(file_path, index=False)
            
    def generate_pdf(self, export_path, content):
        """Generate PDF directly using ReportLab"""
        # Configure the PDF document with proper page setup
        doc = SimpleDocTemplate(
            export_path,
            pagesize=letter,
            rightMargin=36,
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        
        # Calculate available width for content
        available_width = letter[0] - 72  # 72 points = 1 inch (36pt margin on each side)
        
        styles = getSampleStyleSheet()
        story = []
    
        # VI. Course Schedule
        story.append(Paragraph("VI. Calendar", styles['CustomHeading1']))
        
        # Create schedule table with preprocessed data
        schedule_data = [["Date", "Topic", "Readings/Preparation", "Work Due"]]
        
        # Process the schedule entries first to better handle text wrapping
        processed_schedule_data = []
        for entry in self.schedule_entries:
            date = entry['date'].get()
            topic = entry['topic'].get()
            readings = entry['readings'].get("1.0", tk.END).strip()
            work_due = entry['work_due'].get()
            
            # Insert intelligent line breaks for better text wrapping in readings column
            if len(readings) > 80:  # If content is long
                # Add line breaks at natural points
                readings = readings.replace(". ", ".\n")
                readings = readings.replace("; ", ";\n")
                readings = readings.replace("] ", "]\n")
                readings = readings.replace("words)", "words)\n")
            
            # Insert breaks for topic if needed
            if len(topic) > 20:
                topic = topic.replace("; ", ";\n")
                topic = topic.replace(": ", ":\n")
            
            processed_schedule_data.append([date, topic, readings, work_due])
        
        schedule_data.extend(processed_schedule_data)
        
        # Set optimized column widths - adjusted for better balance
        col_widths = [
            0.6*inch,   # Date column (slightly narrower)
            1.1*inch,   # Topic column (slightly narrower)
            3.0*inch,   # Readings column (wider for more content)
            1.0*inch    # Work Due column
        ]
        
        # Create schedule table with the adjusted column widths
        schedule_table = Table(schedule_data, colWidths=col_widths, repeatRows=1)
        
        # Apply table styles with improved word wrapping settings
        schedule_table.setStyle(TableStyle([
            # Header row styling
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            
            # Content rows styling
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 7),  # Smaller font for better fit
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align text to top
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
            
            # Reduced padding for more content space
            ('LEFTPADDING', (0, 0), (-1, -1), 2),
            ('RIGHTPADDING', (0, 0), (-1, -1), 2),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            
            # Critical for word wrapping
            ('WORDWRAP', (0, 0), (-1, -1), True)
        ]))
        
        # Add the table to the story
        story.append(schedule_table)
        
        # Define a page template with page numbers
        def add_page_number(canvas, doc):
            """Add page numbers to each page"""
            canvas.saveState()
            canvas.setFont('Helvetica', 9)
            # Position the page number at the bottom center of the page
            canvas.drawCentredString(letter[0]/2, 20, f"Page {canvas.getPageNumber()}")
            canvas.restoreState()
        
        # Build the document with page numbers
        doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)

    def generate_syllabus(self, export_format="docx"):
        """Generate the final syllabus document"""
        if not self.validate_inputs():
            return

        if export_format == "pdf":
            default_ext = ".pdf"
            file_types = [("PDF Document", "*.pdf"), ("Word Document", "*.docx")]
        else:
            default_ext = ".docx"
            file_types = [("Word Document", "*.docx"), ("PDF Document", "*.pdf")]

        export_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=file_types,
            title="Save Syllabus As"
        )

        if not export_path:
            return

        try:
            if export_path.lower().endswith('.pdf') or export_format == "pdf":
                # 1. Create a temporary .docx file
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
                    docx_path = tmp_docx.name
                doc = self.create_syllabus_document()
                doc.save(docx_path)
                # 2. Convert to PDF
                convert(docx_path, export_path)
                # 3. Remove the temporary .docx
                os.remove(docx_path)
                messagebox.showinfo("Success", f"Syllabus saved as PDF: {export_path}")
            else:
                doc = self.create_syllabus_document()
                doc.save(export_path)
                messagebox.showinfo("Success", f"Syllabus saved as Word document: {export_path}")
        except Exception as e:
            messagebox.showerror("Error", 
                f"Failed to save syllabus:\n{str(e)}\n\n"
                "Please make sure you have write permissions and the file is not open in another program.")
            
    def validate_inputs(self):
        """Validate required inputs before generating syllabus"""
        required_fields = [
            (self.entry_course_num, "Course Number"),
            (self.entry_course_title, "Course Title"),
            (self.entry_term, "Term"),
            (self.entry_credits, "Credits"),
            (self.entry_meeting_times, "Meeting Times"),
            (self.entry_location, "Location"),
            (self.entry_instr_name, "Instructor Name"),
            (self.entry_instr_email, "Instructor Email")
        ]
        
        for field, name in required_fields:
            if not field.get().strip():
                messagebox.showerror("Error", f"{name} is required.")
                return False
        return True

    def create_syllabus_document(self):
        """Create the Word document for the syllabus following the exact format from the example"""
        try:
            doc = Document()
            
            # Add page numbers in the footer
            section = doc.sections[0]
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.text = "Page "
            
            # Add a field for the page number
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            run._element.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._element.append(fldChar2)
            
            # Continue with the rest of the document creation code
            # Title and Course Info (centered)
            title = doc.add_heading(f"{self.entry_course_num.get()}: {self.entry_course_title.get()}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            term = doc.add_paragraph(f"{self.entry_term.get()} ({self.entry_credits.get()} credits)")
            term.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Single space after title
            
            # I. General Information
            doc.add_heading("I. General Information", level=1)
            
            # Meeting times and location - no extra spacing
            p = doc.add_paragraph()
            p.add_run("Meeting days and times: ").bold = True
            p.add_run(self.entry_meeting_times.get())
            
            p = doc.add_paragraph()
            p.add_run("Class location: ").bold = True
            p.add_run(self.entry_location.get())
            
            # Instructor info - compact format
            p = doc.add_paragraph()
            p.add_run("\nInstructor:").bold = True
            
            instructor_info = [
                ("Name:", self.entry_instr_name.get()),
                ("Office:", self.entry_instr_office.get()),
                ("Phone:", self.entry_instr_phone.get()),
                ("Email:", self.entry_instr_email.get()),
                ("Office Hours:", self.entry_instr_office_hours.get())
            ]
            for label, value in instructor_info:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)  # Reduced indentation
                p.paragraph_format.space_after = Pt(0)  # Remove spacing after paragraphs
                p.add_run(f"{label} ").bold = True
                p.add_run(value)
            
            # Teaching Assistants - compact format
            if self.ta_entries:
                p = doc.add_paragraph()
                p.add_run("\nTeaching Assistants:").bold = True
                
                for ta in self.ta_entries:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Name: ").bold = True
                    p.add_run(ta[0].get())
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Email: ").bold = True
                    p.add_run(ta[1].get())
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run("Office Hours: ").bold = True
                    p.add_run(ta[2].get())
            
            # Course Description
            doc.add_heading("Course Description", level=1)
            doc.add_paragraph(self.txt_description.get("1.0", tk.END).strip())
            
            # Prerequisites (moved before General Education)
            doc.add_heading("Prerequisites", level=1)
            doc.add_paragraph(self.entry_prerequisites.get().strip() if hasattr(self, 'entry_prerequisites') else "None")
            
            # --- Add General Education Designation (moved after Prerequisites) ---
            if self.show_gen_ed.get():
                # Define the designation variable
                designation = "Social and Behavioral Sciences (S)"  # Default value
                
                # Add the General Education heading and note
                doc.add_heading(f"General Education Designation: {designation}", level=1)
                
                # Add the full General Education description text
                doc.add_paragraph(gen_ed_default)
                
                # Then continue with the existing code for the success message
                course_num = self.entry_course_num.get()
                doc.add_paragraph(f"Your successful completion of {course_num} with a grade of \"C\" or higher will count towards UF's General Education State Core in {designation}. It will also count towards the State of Florida's Civic Literacy requirement.")
            
            # Course Objectives - Add this section (moved after General Education)
            doc.add_heading("Course Objectives", level=1)
            doc.add_paragraph("All General Education area objectives can be found here.")
            doc.add_paragraph(f"The {self.entry_course_num.get()} curriculum will also cover the following course-specific objectives:")
            
            if hasattr(self, 'objective_entries') and any(obj["entry"].get().strip() for obj in self.objective_entries):
                for i, obj in enumerate(self.objective_entries, 1):
                    obj_text = obj["entry"].get().strip()
                    if obj_text:
                        p = doc.add_paragraph(f"{i}. {obj_text}")
                        p.paragraph_format.left_indent = Inches(0.25)
            
            # Add Student Learning Outcomes
            doc.add_heading("II. Student Learning Outcomes", level=1)
            doc.add_paragraph("A student who successfully completes this course will:")
            
            if hasattr(self, 'outcome_entries') and self.outcome_entries:
                for i, outcome_entry in enumerate(self.outcome_entries, 1):
                    outcome_text = outcome_entry["entry"].get().strip()
                    if outcome_text:
                        p = doc.add_paragraph(f"{i}. {outcome_text}")
                        p.paragraph_format.left_indent = Inches(0.25)
            else:
                # Default outcomes if none provided
                default_outcomes = [
                    "Describe the factual details of the substantive historical episodes under study.",
                    "Identify and analyze foundational developments that shaped history using critical thinking skills.",
                    "Demonstrate an understanding of the primary ideas, values, and perceptions that have shaped history.",
                    "Demonstrate competency in civic literacy."
                ]
                for i, outcome in enumerate(default_outcomes, 1):
                    p = doc.add_paragraph(f"{i}. {outcome}")
                    p.paragraph_format.left_indent = Inches(0.25)
            
            # If General Education is enabled, add the objectives table after the Student Learning Outcomes
            if self.show_gen_ed.get():
                # Add Learning Outcomes table
                doc.add_paragraph()
                doc.add_paragraph(f"Objectives—General Education and {designation}")
                
                # Create table with the correct number of rows and columns
                table_rows = 1  # Header row
                
                # Count rows based on learning objectives entries if they exist
                if hasattr(self, 'learning_objectives_entries') and self.learning_objectives_entries:
                    for category, entries in self.learning_objectives_entries.items():
                        # Skip entries that were removed
                        if 'frame' in entries and not entries['frame'].winfo_exists():
                            continue
                        table_rows += 1
                else:
                    # Default 3 categories if no custom entries
                    table_rows += 3
                    
                table = doc.add_table(rows=table_rows, cols=4)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "CATEGORY"
                
                # Customize the second header based on designation
                slo_header = "SOCIAL SCIENCE SLOS"
                if "Humanities" in designation:
                    slo_header = "HUMANITIES SLOS"
                elif "International" in designation:
                    slo_header = "INTERNATIONAL SLOS"
                elif "Diversity" in designation:
                    slo_header = "DIVERSITY SLOS"
                elif "Biological" in designation:
                    slo_header = "BIOLOGICAL SCIENCES SLOS"
                elif "Physical" in designation:
                    slo_header = "PHYSICAL SCIENCES SLOS"
                elif "Mathematics" in designation:
                    slo_header = "MATHEMATICS SLOS"
                    
                header_cells[1].text = slo_header
                header_cells[2].text = "STATE SLO ASSIGNMENTS"
                header_cells[3].text = "COURSE-SPECIFIC"
                
                # Make header row bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                row_idx = 1
                if hasattr(self, 'learning_objectives_entries') and self.learning_objectives_entries:
                    for category, entries in self.learning_objectives_entries.items():
                        # Skip entries that were removed
                        if 'frame' in entries and not entries['frame'].winfo_exists():
                            continue
                        
                        # For custom categories, use the name_entry value
                        if 'name_entry' in entries:
                            category_name = entries['name_entry'].get()
                        else:
                            category_name = category
                        
                        slo_text = entries['slo'].get("1.0", tk.END).strip()
                        assignments_text = entries['assignments'].get("1.0", tk.END).strip()
                        course_specific_text = entries['course_specific'].get("1.0", tk.END).strip()
                        
                        if row_idx < len(table.rows):
                            row = table.rows[row_idx]
                            row.cells[0].text = category_name
                            row.cells[1].text = slo_text
                            row.cells[2].text = assignments_text
                            row.cells[3].text = course_specific_text
                            row_idx += 1
                else:
                    # Fallback to default table if no custom entries exist
                    default_data = [
                        ("Content", 
                         "Identify, describe, and explain key themes, principles, and terminology; the history, theory and/or methodologies used; and social institutions, structures and processes.", 
                         "Outcomes 1-4\n\nStudents will demonstrate their knowledge of the details of the substantive historical episodes by analyzing primary and secondary sources in short papers, homework assignments, exams, and in-class discussion."),
                        ("Critical Thinking", 
                         "Apply formal and informal qualitative or quantitative analysis effectively to examine the processes and means by which individuals make personal and group decisions. Assess and analyze ethical perspectives in individual and societal decisions.", 
                         "Outcomes 1-4\n\nStudents will demonstrate their ability in applying qualitative and quantitative methods by analyzing primary and secondary sources in short papers, homework assignments, and exams by using critical thinking skills."),
                        ("Communication", 
                         "Communication is the development and expression of ideas in written and oral forms.", 
                         "Outcomes 1-4\n\nStudents will identify and explain key developments that shaped history in written assignments and class discussion.\n\nStudents will demonstrate their understandings of the primary ideas, values, and perceptions that have shaped history and will describe them in written assignments, exams, and class discussion.")
                    ]
                    
                    for i, (category, slo, assignments) in enumerate(default_data):
                        if row_idx < len(table.rows):
                            row = table.rows[row_idx]
                            row.cells[0].text = category
                            row.cells[1].text = slo
                            row.cells[2].text = assignments
                            row.cells[3].text = ""
                            row_idx += 1
            
            # III. Graded Work
            doc.add_heading("III. Graded Work", level=1)
            
            # Materials (if provided)
            if hasattr(self, 'materials_text') and self.materials_text.get("1.0", tk.END).strip():
                doc.add_heading("Required Materials", level=2)
                materials_text = self.materials_text.get("1.0", tk.END).strip()
                # --- Use markup parser for formatted output ---
                self.parse_materials_markup(materials_text, doc=doc)
                # Always include the Materials Fee value from the fee_entry
                fee_value = ""
                if hasattr(self, 'fee_entry') and self.fee_entry.get().strip():
                    fee_value = self.fee_entry.get().strip()
                else:
                    fee_value = "0.00"
                p = doc.add_paragraph()
                p.add_run("\nMaterials Fee: $").bold = True
                p.add_run(fee_value)

            # ==== CONTINUING AFTER REQUIRED MATERIALS SECTION ====
            
            # Grading Components (Categories and Assignments)
            if hasattr(self, 'category_frames') and self.category_frames:
                doc.add_heading("Grading Components", level=2)
                
                # Create a table for categories and weights
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Category"
                header_cells[1].text = "Weight"
                
                # Make header row bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for category in self.category_frames:
                    name = category["name"].get().strip()
                    weight = category["weight"].get().strip()
                    
                    if name and weight:
                        row_cells = table.add_row().cells
                        row_cells[0].text = name
                        row_cells[1].text = f"{weight}%"
                
                # Category descriptions and assignments
                for category in self.category_frames:
                    name = category["name"].get().strip()
                    desc = category["description"].get("1.0", tk.END).strip()
                    
                    if name and desc:
                        p = doc.add_paragraph()
                        p.add_run(f"\n{name}: ").bold = True
                        p.add_run(desc)
                    
                    # Assignments for this category
                    if "assignments" in category and category["assignments"]:
                        has_assignments = False
                        for assignment in category["assignments"]:
                            title = assignment["title"].get().strip()
                            due_date = assignment["due date"].get().strip()
                            points = assignment["points"].get().strip()
                            description = assignment["description"].get("1.0", tk.END).strip() if "description" in assignment else ""
                            
                            if title:
                                if not has_assignments:
                                    p = doc.add_paragraph()
                                    p.add_run(f"{name} Assignments:").bold = True
                                    has_assignments = True
                                
                                p = doc.add_paragraph()
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.add_run(f"• {title}")
                                
                                if due_date:
                                    p.add_run(f" (Due: {due_date})")
                                if points:
                                    p.add_run(f" - {points} points")
                                
                                if description:
                                    p = doc.add_paragraph(description)
                                    p.paragraph_format.left_indent = Inches(0.5)
            
            # Grading Scale
            doc.add_heading("Grading Scale", level=2)
            
            # Create table for grading scale
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Letter Grade"
            header_cells[1].text = "Number Grade"
            
            # Make header row bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Grade scale data
            grades = [
                ("A", "100-93"),
                ("A-", "92-90"),
                ("B+", "89-87"),
                ("B", "86-83"),
                ("B-", "82-80"),
                ("C+", "79-77"),
                ("C", "76-73"),
                ("C-", "72-70"),
                ("D+", "69-67"),
                ("D", "66-63"),
                ("D-", "62-60"),
                ("E", "59-0")
            ]
            
            for letter, number in grades:
                row_cells = table.add_row().cells
                row_cells[0].text = letter
                row_cells[1].text = number
            
            # Add UF grading policies note
            p = doc.add_paragraph()
            p.add_run("See the UF Catalog's \"Grades and Grading Policies\" for information on how UF assigns grade points.")
            
            # Add minimum grade note
            p = doc.add_paragraph()
            p.add_run("Note: A minimum grade of C is required to earn General Education credit.")
            
            # Instructions for Submitting Written Assignments
            doc.add_heading("Instructions for Submitting Written Assignments", level=1)
            doc.add_paragraph("All written assignments must be submitted as Word documents (.doc or .docx) through the \"Assignments\" portal in Canvas by the specified deadlines. Do NOT send assignments as PDF files.")
            
            # Add University Assessment Policy if checked
            if hasattr(self, 'assessment_policy_var') and self.assessment_policy_var.get():
                doc.add_heading("University Assessment Policies", level=2)
                doc.add_paragraph("Requirements for make-up exams, assignments, and other work in this course are consistent with university policies that can be found in the Catalog.")
            
            # Add Extensions & Make-Up Exams if checked
            if hasattr(self, 'extensions_policy_var') and self.extensions_policy_var.get():
                doc.add_heading("Extensions & Make-Up Exams", level=2)
                doc.add_paragraph("Only the professor can authorize an extension or make-up exam, and all requests must be supported by documentation from a medical provider, Student Health Services, the Disability Resource Center, or the Dean of Students Office. Requirements for attendance and make-up exams, assignments, and other work in this course are consistent with university policies: https://catalog.ufl.edu/ugrad/current/regulations/info/attendance.aspx")
            
            # Add Late Submissions policy
            doc.add_heading("Late Submissions", level=2)
            if hasattr(self, 'late_policy_text') and self.late_policy_text.get("1.0", tk.END).strip():
                doc.add_paragraph(self.late_policy_text.get("1.0", tk.END).strip())
            elif hasattr(self, 'late_policy_var') and self.late_policy_var.get() in self.late_policies:
                selected_policy = self.late_policy_var.get()
                doc.add_paragraph(self.late_policies[selected_policy])
            else:
                doc.add_paragraph("Late submission policy not specified.")
            
            # Add Extra Credit policy
            doc.add_heading("Extra Credit", level=2)
            if hasattr(self, 'extra_credit_text') and self.extra_credit_text.get("1.0", tk.END).strip():
                doc.add_paragraph(self.extra_credit_text.get("1.0", tk.END).strip())
            elif hasattr(self, 'extra_credit_var') and self.extra_credit_var.get() in self.extra_credit_policies:
                selected_policy = self.extra_credit_var.get()
                doc.add_paragraph(self.extra_credit_policies[selected_policy])
            else:
                doc.add_paragraph("Extra credit policy not specified.")
            
            # Canvas Policy (always included)
            doc.add_heading("Canvas", level=2)
            if hasattr(self, 'canvas_policy_text') and self.canvas_policy_text.get("1.0", tk.END).strip() != "-Replace with your Canvas Policies-":
                doc.add_paragraph(self.canvas_policy_text.get("1.0", tk.END).strip())
            else:
                doc.add_paragraph(canvas_policy_default)
            
            # Technology Policy
            doc.add_heading("Technology in the Classroom", level=2)
            if hasattr(self, 'technology_policy_text') and self.technology_policy_text.get("1.0", tk.END).strip() != "-Replace with your Technology Policies-":
                doc.add_paragraph(self.technology_policy_text.get("1.0", tk.END).strip())
            else:
                doc.add_paragraph(technology_policy_default)
            
            # Communication Policy
            doc.add_heading("Class Communication Policy", level=2)
            if hasattr(self, 'communication_policy_text'):
                doc.add_paragraph(self.communication_policy_text.get("1.0", tk.END).strip())
            
            # Add Assignment Support section if enabled
            if self.outside_support_var.get():
                doc.add_heading("Assignment Support Outside the Classroom", level=2)
                if hasattr(self, 'support_text'):
                    doc.add_paragraph(self.support_text.get("1.0", tk.END).strip())
                else:
                    doc.add_paragraph("You are welcome to come to regular office hours or to schedule an individual appointment with your professor or TA. When needed, I also encourage you to seek support from the academic resources listed on this syllabus.")
            
            # IV. Evaluations section
            doc.add_heading("IV. Evaluations", level=1)
            eval_text = (
                "UF course evaluation process\n"
                "Students are expected to provide professional and respectful feedback on the quality of "
                "instruction in this course by completing course evaluations online. Students can complete "
                "evaluations in three ways:\n\n"
                "1. The email they receive from GatorEvals,\n"
                "2. Their Canvas course menu under GatorEvals, or\n"
                "3. The central portal at https://my-ufl.bluera.com\n\n"
                "Guidance on how to provide constructive feedback is available at "
                "https://gatorevals.aa.ufl.edu/students/. Students will be notified when the evaluation "
                "period opens. Summaries of course evaluation results are available to students at "
                "https://gatorevals.aa.ufl.edu/public-results/."
            )
            doc.add_paragraph(eval_text)
            
            # V. University Policies and Resources
            doc.add_heading("V. University Policies and Resources", level=1)
            
            # Accommodations policy
            doc.add_heading("Students requiring accommodation", level=2)
            doc.add_paragraph(accommodations_default)
            
            # University Honesty Policy
            doc.add_heading("University Honesty Policy", level=2)
            doc.add_paragraph(honesty_plagiarism_default)
            
            # In-class recording policy (if enabled)
            if hasattr(self, 'in_class_recording_var') and self.in_class_recording_var.get():
                doc.add_heading("In-class recording", level=2)
                doc.add_paragraph(recording_policy_default)
            
            # Conflict resolution (if enabled)
            if hasattr(self, 'conflict_resolution_var') and self.conflict_resolution_var.get():
                doc.add_heading("Procedure for conflict resolution", level=2)
                doc.add_paragraph(conflict_resolution_default)
            
            # Campus Resources (if enabled)
            if self.campus_resources_var.get():
                doc.add_heading("Campus Resources", level=2)
                doc.add_paragraph(campus_resources_default)
            
            # Academic Resources (if enabled)
            if self.academic_resources_var.get():
                doc.add_heading("Academic Resources", level=2)
                doc.add_paragraph(academic_resources_default)
            
            # VI. Course Schedule (Calendar)
            doc.add_heading("VI. Calendar", level=1)
            
            # Create schedule table if there are entries
            if hasattr(self, 'schedule_entries') and self.schedule_entries:
                # Create table with headers
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Set column headers
                header_cells = table.rows[0].cells
                header_cells[0].text = "Date"
                header_cells[1].text = "Topic"
                header_cells[2].text = "Readings/Preparation"
                header_cells[3].text = "Work Due"
                
                # Make header row bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add each schedule entry as a row
                for entry in self.schedule_entries:
                    date_text = entry["date"].get().strip()
                    topic_text = entry["topic"].get().strip()
                    readings_text = entry["readings"].get("1.0", tk.END).strip()
                    work_due_text = entry["work_due"].get().strip()
                    
                    # Skip empty rows
                    if not any([date_text, topic_text, readings_text, work_due_text]):
                        continue
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = date_text
                    row_cells[1].text = topic_text
                    row_cells[2].text = readings_text
                    row_cells[3].text = work_due_text
            else:
                doc.add_paragraph("Schedule will be provided separately.")

            return doc
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while generating the document: {e}")
            import traceback
            traceback.print_exc()
            # Return an empty document so we don't crash
            return Document()

    def run(self):
        """Start the application"""
        self.root.mainloop()

    def add_category(self, name="", weight="", description=""):
        """Add a new empty grading category"""
        frame = ttk.LabelFrame(self.categories_frame)
        frame.pack(fill=tk.X, pady=5)
        
        # Category header
        header_frame = ttk.Frame(frame)
        header_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(header_frame, text="Category Name:").pack(side=tk.LEFT)
        name_entry = ttk.Entry(header_frame, width=30)
        name_entry.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(header_frame, text="Weight (%):").pack(side=tk.LEFT)
        weight_entry = ttk.Entry(header_frame, width=5)
        weight_entry.pack(side=tk.LEFT, padx=5)
        
        # Description
        ttk.Label(frame, text="Description:").pack(anchor="w", padx=5)
        desc_text = scrolledtext.ScrolledText(frame, width=60, height=4, wrap=tk.WORD)
        desc_text.pack(fill=tk.X, padx=5, pady=5)
        self.add_mousewheel_scrolling(desc_text)
        
        # Assignments section
        assignments_frame = ttk.Frame(frame)
        assignments_frame.pack(fill=tk.X, padx=5, pady=5)
        
        assignments = []
        
        def add_assignment():
            assignment_frame = ttk.Frame(assignments_frame)
            assignment_frame.pack(fill=tk.X, pady=2)
            
            ttk.Label(assignment_frame, text="Title:").pack(side=tk.LEFT)
            title_entry = ttk.Entry(assignment_frame, width=30)
            title_entry.pack(side=tk.LEFT, padx=2)
            
            ttk.Label(assignment_frame, text="Due:").pack(side=tk.LEFT)
            due_entry = ttk.Entry(assignment_frame, width=15)
            due_entry.pack(side=tk.LEFT, padx=2)
            
            ttk.Label(assignment_frame, text="Points:").pack(side=tk.LEFT)
            points_entry = ttk.Entry(assignment_frame, width=5)
            points_entry.pack(side=tk.LEFT, padx=2)
            
            # Description for the assignment
            ttk.Label(assignment_frame, text="Description:").pack(side=tk.LEFT)
            description_text = scrolledtext.ScrolledText(assignment_frame, width=40, height=3, wrap=tk.WORD)
            description_text.pack(side=tk.LEFT, padx=2)
            
            def remove_assignment():
                assignment_frame.destroy()
                assignments.remove(assignment_dict)
            
            remove_btn = ttk.Button(assignment_frame, text="×", 
                                  command=remove_assignment,
                                  style="Delete.TButton")
            remove_btn.pack(side=tk.LEFT, padx=2)
            
            assignment_dict = {
                "frame": assignment_frame,
                "title": title_entry,
                "due date": due_entry,
                "points": points_entry,
                "description": description_text  # Add description to the dictionary
            }
            assignments.append(assignment_dict)
        
        ttk.Button(frame, text="Add Assignment", command=add_assignment).pack(anchor="w", padx=5, pady=5)
        
        def remove_category():
            frame.destroy()
            self.category_frames.remove(category_dict)
        
        ttk.Button(frame, text="Remove Category", command=remove_category).pack(anchor="w", padx=5, pady=5)
        
        category_dict = {
            "frame": frame,
            "name": name_entry,
            "weight": weight_entry,
            "description": desc_text,
            "assignments": assignments
        }
        
        self.category_frames.append(category_dict)
        return category_dict

    def gather_content(self):
        """Gather all content from form fields"""
        try:
            content = {
                "course_info": {
                    "course_num": self.entry_course_num.get() if hasattr(self, 'entry_course_num') else "",
                    "course_title": self.entry_course_title.get() if hasattr(self, 'entry_course_title') else "",
                    "term": self.entry_term.get() if hasattr(self, 'entry_term') else "",
                    "credits": self.entry_credits.get() if hasattr(self, 'entry_credits') else "",
                    # Add prerequisites here
                    "prerequisites": self.entry_prerequisites.get() if hasattr(self, 'entry_prerequisites') else "",
                    "meeting_times": self.entry_meeting_times.get() if hasattr(self, 'entry_meeting_times') else "",
                    "location": self.entry_location.get() if hasattr(self, 'entry_location') else "",
                    "description": self.txt_description.get("1.0", tk.END).strip() if hasattr(self, 'txt_description') else "",
                    "objectives": "\n".join([obj["entry"].get().strip() for obj in self.objective_entries if obj["entry"].get().strip()]) if hasattr(self, 'objective_entries') else ""
                },
                "instructor_info": {
                    "name": self.entry_instr_name.get() if hasattr(self, 'entry_instr_name') else "",
                    "office": self.entry_instr_office.get() if hasattr(self, 'entry_instr_office') else "",
                    "phone": self.entry_instr_phone.get() if hasattr(self, 'entry_instr_phone') else "",
                    "email": self.entry_instr_email.get() if hasattr(self, 'entry_instr_email') else "",
                    "office_hours": self.entry_instr_office_hours.get() if hasattr(self, 'entry_instr_office_hours') else ""
                },
                "tas": [],
                "outcomes": [],
                # Removed gen_ed section as it's handled by show_gen_ed boolean
                "schedule": [],
                "grading_categories": [],
                "optional_policies": {}, # Will be populated below
                "late_policy": "",
                "extra_credit_policy": "",
                "canvas_policy": self.canvas_policy_text.get("1.0", tk.END).strip() if hasattr(self, 'canvas_policy_text') else "-Replace with your Canvas Policies-",
                "technology_policy": self.technology_policy_text.get("1.0", tk.END).strip() if hasattr(self, 'technology_policy_text') else "-Replace with your Technology Policies-",
                "communication_policy": self.communication_policy_text.get("1.0", tk.END).strip() if hasattr(self, 'communication_policy_text') else "-Replace with your policy of communication with students-",
                "support_policy": self.support_text.get("1.0", tk.END).strip() if hasattr(self, 'support_text') else "-Write your policy for assisting students outside of the classroom here-",
                "learning_objectives": {}
            }

            # Gather outcomes from the new numbered list
            if hasattr(self, 'outcome_entries'):
                content["outcomes"] = [
                    {"text": outcome_entry["entry"].get().strip()}
                    for outcome_entry in self.outcome_entries if outcome_entry["entry"].get().strip()
                ]

            # Gather TAs
            if hasattr(self, 'ta_entries'):
                for ta_entry_widgets in self.ta_entries:
                    if len(ta_entry_widgets) == 3: # Ensure it has name, email, hours
                         content["tas"].append({
                            "name": ta_entry_widgets[0].get(),
                            "email": ta_entry_widgets[1].get(),
                            "office_hours": ta_entry_widgets[2].get()
                        })

            # Gather schedule
            if hasattr(self, 'schedule_entries'):
                for entry in self.schedule_entries:
                    content["schedule"].append({
                        "date": entry["date"].get(),
                        "topic": entry["topic"].get(),
                        "readings": entry["readings"].get("1.0", tk.END).strip(),
                        "work_due": entry["work_due"].get()
                    })

            # Gather assignment categories
            if hasattr(self, 'category_frames'):
                for category in self.category_frames:
                    assignments_data = []
                    if "assignments" in category:
                        for assignment in category["assignments"]:
                            assignments_data.append({
                                "title": assignment["title"].get(),
                                "due_date": assignment["due date"].get(),
                                "points": assignment["points"].get(),
                                "description": assignment["description"].get("1.0", tk.END).strip()
                            })
                    content["grading_categories"].append({
                        "name": category["name"].get(),
                        "weight": category["weight"].get(),
                        "description": category["description"].get("1.0", tk.END).strip(),
                        "assignments": assignments_data
                    })

            # Gather learning objectives table data
            if hasattr(self, 'learning_objectives_entries'):
                for category_key, entries in self.learning_objectives_entries.items():
                     # Use name_entry if it exists (for custom categories), otherwise use the key
                    category_name = entries.get('name_entry').get() if 'name_entry' in entries else category_key
                    if category_name: # Only add if category name is not empty
                        content["learning_objectives"][category_name] = {
                            "slo": entries['slo'].get("1.0", tk.END).strip(),
                            "assignments": entries['assignments'].get("1.0", tk.END).strip(),
                            "course_specific": entries['course_specific'].get("1.0", tk.END).strip()
                        }

            # Gather optional policies boolean values
            if hasattr(self, 'optional_policies'):
                for policy_name, policy_var in self.optional_policies.items():
                    content["optional_policies"][policy_name] = policy_var.get()
            # Also gather the state of the main checkboxes
            content["optional_policies"]["assessment"] = self.assessment_policy_var.get() if hasattr(self, 'assessment_policy_var') else True
            content["optional_policies"]["extensions"] = self.extensions_policy_var.get() if hasattr(self, 'extensions_policy_var') else True
            content["optional_policies"]["outside_support"] = self.outside_support_var.get() if hasattr(self, 'outside_support_var') else True
            content["optional_policies"]["show_gen_ed"] = self.show_gen_ed.get() if hasattr(self, 'show_gen_ed') else True


            if hasattr(self, 'materials_text'):
                content["materials"] = {
                    "required": self.materials_text.get("1.0", tk.END).strip(),
                    "fee": self.fee_entry.get() if hasattr(self, 'fee_entry') else ""
                }

            # Add policy details if available
            if hasattr(self, 'late_policy_var'):
                content["late_policy"] = self.late_policy_var.get()
                # Store the actual text if it's custom or selected
                if hasattr(self, 'late_policy_text'):
                     content["late_policy_text"] = self.late_policy_text.get("1.0", tk.END).strip()


            if hasattr(self, 'extra_credit_var'):
                content["extra_credit_policy"] = self.extra_credit_var.get()
                # Store the actual text if it's custom or selected
                if hasattr(self, 'extra_credit_text'):
                    content["extra_credit_policy_text"] = self.extra_credit_text.get("1.0", tk.END).strip()


            return content

        except Exception as e:
            print(f"Error gathering content: {e}")
            traceback.print_exc()
            # Return a minimal content structure to prevent errors
            return {
                "course_info": {
                    "course_num": "", "course_title": "", "term": "", "credits": "", "prerequisites": "",
                    "meeting_times": "", "location": "", "description": "", "objectives": ""
                },
                "instructor_info": {"name": "", "office": "", "phone": "", "email": "", "office_hours": ""},
                "tas": [], "outcomes": [], "schedule": [], "grading_categories": [],
                "optional_policies": {}, "learning_objectives": {}, "materials": {"required": "", "fee": ""},
                "late_policy": "", "extra_credit_policy": "", "canvas_policy": "", "technology_policy": "",
                "communication_policy": "", "support_policy": ""
            }

    def create_document_preview_tab(self):
        """Create a tab for previewing the entire document as it will appear in final form"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Document Preview")
        self.add_mousewheel_scrolling(tab, tab)
        
        # Create a scrollable preview area
        preview_frame = ttk.Frame(tab)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Controls at the top
        control_frame = ttk.Frame(preview_frame)
        control_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(control_frame, text="Document Preview", style="Heading.TLabel").pack(side=tk.LEFT, padx=5)
        
        refresh_btn = ttk.Button(control_frame, text="Refresh Preview", 
                              command=self.update_document_preview,
                              style="Action.TButton")
        refresh_btn.pack(side=tk.RIGHT, padx=5)
        
        # Create a canvas with scrollbar for the preview content
        canvas_frame = ttk.Frame(preview_frame, relief=tk.SUNKEN, borderwidth=1)
        canvas_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.preview_canvas = tk.Canvas(canvas_frame, bg="white")
        scrollbar_y = ttk.Scrollbar(canvas_frame, orient="vertical", command=self.preview_canvas.yview)
        scrollbar_x = ttk.Scrollbar(canvas_frame, orient="horizontal", command=self.preview_canvas.xview)
        
        self.preview_content_frame = ttk.Frame(self.preview_canvas, style="Preview.TFrame")
        self.preview_content_frame.bind(
            "<Configure>",
            lambda e: self.preview_canvas.configure(
                scrollregion=self.preview_canvas.bbox("all")
            )
        )
        self.preview_canvas.create_window((0, 0), window=self.preview_content_frame, anchor="nw")
        self.preview_canvas.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)
        
        # Pack the scrollbars and canvas
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
        self.preview_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Add mousewheel scrolling
        self.add_mousewheel_scrolling(self.preview_canvas, self.preview_canvas)
        self.add_mousewheel_scrolling(self.preview_content_frame, self.preview_canvas)
        
        # Bind mousewheel events to the preview content so the canvas scrolls
        self.preview_canvas.bind("<Enter>", lambda e: self.preview_canvas.focus_set())
        self.preview_canvas.bind("<MouseWheel>", lambda event: self.preview_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units"))
        
        # Initial preview generation
        self.update_document_preview()

    def update_document_preview(self):
        """Update the document preview with current content"""
        # Clear previous content
        for widget in self.preview_content_frame.winfo_children():
            widget.destroy()

        try:
            # Create container for all preview sections
            content_container = ttk.Frame(self.preview_content_frame, style="Preview.TFrame") # Use Preview style
            content_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

            # Title and Course Info
            title_text = f"{self.entry_course_num.get()}: {self.entry_course_title.get()}"
            title = ttk.Label(content_container, text=title_text,
                           font=("Times New Roman", 14, "bold"),
                           style="Preview.TLabel") # Use Preview style
            title.pack(pady=5)

            term_text = f"{self.entry_term.get()} ({self.entry_credits.get()} credits)"
            term = ttk.Label(content_container, text=term_text,
                          font=("Times New Roman", 12),
                          style="Preview.TLabel") # Use Preview style
            term.pack(pady=2)

            ttk.Separator(content_container).pack(fill=tk.X, pady=10)

            # I. General Information
            self._add_preview_section(content_container, "I. General Information", 12, "bold")

            self._add_preview_field(content_container, "Meeting days and times:", self.entry_meeting_times.get())
            self._add_preview_field(content_container, "Class location:", self.entry_location.get())

            # Instructor info
            self._add_preview_text(content_container, "\nInstructor:", bold=True)

            instructor_frame = ttk.Frame(content_container, style="Preview.TFrame")
            instructor_frame.pack(fill=tk.X, padx=20, pady=2)

            instructor_info = [
                ("Name:", self.entry_instr_name.get()),
                ("Office:", self.entry_instr_office.get()),
                ("Phone:", self.entry_instr_phone.get()),
                ("Email:", self.entry_instr_email.get()),
                ("Office Hours:", self.entry_instr_office_hours.get())
            ]

            for i, (label, value) in enumerate(instructor_info):
                 self._add_preview_field(instructor_frame, label, value, indent=0) # No extra indent needed here

            # Teaching Assistants
            if self.ta_entries:
                 self._add_preview_text(content_container, "\nTeaching Assistants:", bold=True)
                 ta_frame = ttk.Frame(content_container, style="Preview.TFrame")
                 ta_frame.pack(fill=tk.X, padx=20, pady=2)
                 for ta_entry_widgets in self.ta_entries:
                     if len(ta_entry_widgets) == 3:
                         self._add_preview_field(ta_frame, "Name:", ta_entry_widgets[0].get(), indent=0)
                         self._add_preview_field(ta_frame, "Email:", ta_entry_widgets[1].get(), indent=0)
                         self._add_preview_field(ta_frame, "Office Hours:", ta_entry_widgets[2].get(), indent=0)
                         ttk.Separator(ta_frame).pack(fill=tk.X, pady=2) # Add separator between TAs


            # Course Description
            self._add_preview_section(content_container, "Course Description", 12, "bold")
            self._add_preview_text(content_container, self.txt_description.get("1.0", tk.END).strip())

            # Prerequisites
            self._add_preview_section(content_container, "Prerequisites", 12, "bold")
            prerequisites = self.entry_prerequisites.get().strip() if hasattr(self, 'entry_prerequisites') else "None"
            self._add_preview_text(content_container, prerequisites or "None")
            
            # Gen Ed (if applicable)
            if self.show_gen_ed.get():
                # Title bolded
                self._add_preview_section(content_container, "General Education Designation: Social and Behavioral Sciences (S)", 12, "bold")
                # Description text
                self._add_preview_text(content_container, gen_ed_default)
            
            # Course Objectives
            self._add_preview_section(content_container, "Course Objectives", 12, "bold")
            objectives_frame = ttk.Frame(content_container, style="Preview.TFrame")
            objectives_frame.pack(fill=tk.X, padx=10, pady=2)
            for i, obj in enumerate(self.objective_entries, 1):
                obj_text = obj["entry"].get().strip()
                if obj_text:
                    self._add_preview_text(objectives_frame, f"{i}. {obj_text}", indent=10)
            
            # II. Student Learning Outcomes
            self._add_preview_section(content_container, "II. Student Learning Outcomes", 12, "bold")
            self._add_preview_text(content_container, "A student who successfully completes this course will:")
            
            outcomes_frame = ttk.Frame(content_container, style="Preview.TFrame")
            outcomes_frame.pack(fill=tk.X, padx=10, pady=2)
            
            if hasattr(self, 'outcome_entries') and self.outcome_entries:
                # Use the entries from the new numbered list UI
                for i, outcome_entry in enumerate(self.outcome_entries, 1):
                    outcome_text = outcome_entry["entry"].get().strip()
                    if outcome_text:
                        self._add_preview_text(outcomes_frame, f"{i}. {outcome_text}", indent=10)
            else:
                # Default outcomes
                default_outcomes = [
                    "Describe the factual details of the substantive historical episodes under study.",
                    "Identify and analyze foundational developments that shaped history using critical thinking skills.",
                    "Demonstrate an understanding of the primary ideas, values, and perceptions that have shaped history.",
                    "Demonstrate competency in civic literacy."
                ]
                for i, outcome in enumerate(default_outcomes, 1):
                    self._add_preview_text(outcomes_frame, f"{i}. {outcome}", indent=10)
            
            # Learning Objectives Table
            if hasattr(self, 'learning_objectives_entries') and self.learning_objectives_entries:
                # Get the current gen ed designation
                designation = "Social and Behavioral Sciences (S)"  # Default
                if hasattr(self, 'gen_ed_designation'):
                    designation = self.gen_ed_designation.get()
                
                title_text = f"Objectives—General Education and {designation}"
                self._add_preview_text(content_container, title_text, bold=True, pady=10)
                
                # Create table
                table_frame = ttk.Frame(content_container, borderwidth=1, relief=tk.SOLID)
                table_frame.pack(fill=tk.X, pady=5, padx=10)
                
                # Table headers (now 4 columns)
                slo_header = "SOCIAL SCIENCE SLOS"
                if "Humanities" in designation:
                    slo_header = "HUMANITIES SLOS"
                elif "International" in designation:
                    slo_header = "INTERNATIONAL SLOS"
                elif "Diversity" in designation:
                    slo_header = "DIVERSITY SLOS"
                elif "Biological" in designation:
                    slo_header = "BIOLOGICAL SCIENCES SLOS"
                elif "Physical" in designation:
                    slo_header = "PHYSICAL SCIENCES SLOS"
                elif "Mathematics" in designation:
                    slo_header = "MATHEMATICS SLOS"
                
                headers = ["CATEGORY", slo_header, "STATE SLO ASSIGNMENTS", "COURSE-SPECIFIC"]
                header_row = ttk.Frame(table_frame, style="Preview.TFrame")
                header_row.pack(fill=tk.X)
                
                for i, text in enumerate(headers):
                    cell = ttk.Label(header_row, text=text, background="#808080", foreground="white",
                                  font=("Arial", 9, "bold"), padding=5)
                    cell.grid(row=0, column=i, sticky="nsew", padx=1, pady=1)
                    header_row.columnconfigure(i, weight=1)
                
                # Table content rows
                row_idx = 0
                for category, entries in self.learning_objectives_entries.items():
                    # Skip entries that were removed
                    if 'frame' in entries and not entries['frame'].winfo_exists():
                        continue
                    
                    row_frame = ttk.Frame(table_frame, style="Preview.TFrame")
                    row_frame.pack(fill=tk.X)
                    
                    # For custom categories, use the name_entry value
                    if 'name_entry' in entries:
                        category_text = entries['name_entry'].get()
                    else:
                        category_text = category
                    
                    slo_text = entries['slo'].get("1.0", tk.END).strip()
                    assignments_text = entries['assignments'].get("1.0", tk.END).strip()
                    course_specific_text = entries['course_specific'].get("1.0", tk.END).strip()
                    
                    # Create table cells with wider wraplength for course-specific
                    ttk.Label(row_frame, text=category_text, background="white",
                           wraplength=150, padding=5, borderwidth=1, relief=tk.SOLID).grid(
                        row=row_idx, column=0, sticky="nsew")
                    ttk.Label(row_frame, text=slo_text, background="white",
                           wraplength=220, padding=5, borderwidth=1, relief=tk.SOLID).grid(
                        row=row_idx, column=1, sticky="nsew")
                    ttk.Label(row_frame, text=assignments_text, background="white",
                           wraplength=220, padding=5, borderwidth=1, relief=tk.SOLID).grid(
                        row=row_idx, column=2, sticky="nsew")
                    ttk.Label(row_frame, text=course_specific_text, background="white",
                           wraplength=350, padding=5, borderwidth=1, relief=tk.SOLID).grid(
                        row=row_idx, column=3, sticky="nsew")
                    
                    for i in range(4):
                        row_frame.columnconfigure(i, weight=1)
                    
                    row_idx += 1
                # If no entries were added/displayed,
                if row_idx == 0:
                    self._add_preview_text(content_container, "Learning objectives table is empty.")
            
            # III. Graded Work
            self._add_preview_section(content_container, "III. Graded Work", 12, "bold")
            
            # Required Materials
            if hasattr(self, 'materials_text') and self.materials_text.get("1.0", tk.END).strip():
                self._add_preview_section(content_container, "Required Materials", 11, "bold")
                materials = self.materials_text.get("1.0", tk.END).strip()
                self._add_preview_text(content_container, materials)
                # Show Materials Fee if present
                fee_value = ""
                if hasattr(self, 'fee_entry') and self.fee_entry.get().strip():
                    fee_value = self.fee_entry.get().strip()
                else:
                    fee_value = "0.00"
                self._add_preview_text(content_container, f"Materials Fee: ${fee_value}", bold=True)
            
            # Grading Components
            if hasattr(self, 'category_frames') and self.category_frames:
                self._add_preview_section(content_container, "Grading Components", 11, "bold")
                
                # Create table for grade components
                grade_frame = ttk.Frame(content_container, style="Preview.TFrame")
                grade_frame.pack(fill=tk.X, padx=10, pady=5)
                
                # Header for grade table
                ttk.Label(grade_frame, text="Category", font=("Arial", 10, "bold"), width=25).grid(
                    row=0, column=0, sticky="w", padx=5, pady=2)
                ttk.Label(grade_frame, text="Weight", font=("Arial", 10, "bold"), width=10).grid(
                    row=0, column=1, sticky="w", padx=5, pady=2)
                ttk.Separator(grade_frame, orient=tk.HORIZONTAL).grid(
                    row=1, column=0, columnspan=2, sticky="ew", pady=2)
                
                # Add each category
                row = 2
                for i, category in enumerate(self.category_frames):
                    if category["name"].get().strip() and category["weight"].get().strip():
                        ttk.Label(grade_frame, text=category["name"].get()).grid(
                            row=row, column=0, sticky="w", padx=5, pady=2)
                        ttk.Label(grade_frame, text=f"{category['weight'].get()}%").grid(
                            row=row, column=1, sticky="w", padx=5, pady=2)
                        row += 1
                
                # Check for descriptions and assignments
                for category in self.category_frames:
                    if "description" in category and category["description"].get("1.0", tk.END).strip():
                        self._add_preview_text(content_container, f"\n{category['name'].get()}:", bold=True)
                        self._add_preview_text(content_container, 
                                           category["description"].get("1.0", tk.END).strip(), indent=10)
                    
                    # Add assignments for this category if any
                    if "assignments" in category and category["assignments"]:
                        has_assignments = False
                        for assignment in category["assignments"]:
                            if assignment["title"].get().strip():
                                if not has_assignments:
                                    self._add_preview_text(content_container, f"\n{category['name'].get()} Assignments:", bold=True)
                                    has_assignments = True
                                assignment_text = f"• {assignment['title'].get()}"
                                if assignment["due date"].get().strip():
                                    assignment_text += f" (Due: {assignment['due date'].get()})"
                                if assignment["points"].get().strip():
                                    assignment_text += f" - {assignment['points'].get()} points"
                                self._add_preview_text(content_container, assignment_text, indent=10)
                                # Add assignment description if available
                                if "description" in assignment and assignment["description"].get("1.0", tk.END).strip():
                                    self._add_preview_text(content_container, assignment["description"].get("1.0", tk.END).strip(), indent=20)
            
            # Add more sections as needed
            # Grading Scale
            self._add_preview_section(content_container, "Grading Scale", 11, "bold")
            
            # Create grading scale data
            grades = [
                ["Letter Grade", "Number Grade"],
                ["A", "100-93"],
                ["A-", "92-90"],
                ["B+", "89-87"],
                ["B", "86-83"],
                ["B-", "82-80"],
                ["C+", "79-77"],
                ["C", "76-73"],
                ["C-", "72-70"],
                ["D+", "69-67"],
                ["D", "66-63"],
                ["D-", "62-60"],
                ["E", "59-0"]
            ]
            
            # Create table frame
            table_frame = ttk.Frame(content_container, borderwidth=1, relief=tk.SOLID)
            table_frame.pack(fill=tk.X, pady=5, padx=10)
            
            # Table headers
            header_row = ttk.Frame(table_frame, style="Preview.TFrame")
            header_row.pack(fill=tk.X)
            
            for i, header in enumerate(grades[0]):
                cell = ttk.Label(header_row, text=header, background="#808080", foreground="white",
                              font=("Arial", 9, "bold"), padding=5)
                cell.grid(row=0, column=i, sticky="nsew", padx=1, pady=1)
                header_row.columnconfigure(i, weight=1)
            
            # Table content rows
            for i, grade in enumerate(grades[1:]):
                data_row = ttk.Frame(table_frame, style="Preview.TFrame")
                data_row.pack(fill=tk.X)
                
                for j, value in enumerate(grade):
                    cell = ttk.Label(data_row, text=value, background="white",
                                  padding=5, borderwidth=1, relief=tk.SOLID)
                    cell.grid(row=0, column=j, sticky="nsew", padx=1, pady=1)
                    data_row.columnconfigure(j, weight=1)
            
            
            # Add the grading policy text with a hyperlink
            policy_text = "See the UF Catalog's \"Grades and Grading Policies\" for information on how UF assigns grade points."
            link = "https://catalog.ufl.edu/UGRD/academic-regulations/grades-grading-policies/"
            self._add_preview_text(content_container, policy_text)
            
            # Add the note about minimum grades
            note_text = "Note: A minimum grade of C is required to earn General Education credit."
            self._add_preview_text(content_container, note_text)
            
            # University Policies section
            #self._add_preview_section(content_container, "University Policies", 12, "bold")
            
            # Add University Assessment Policy if checked
            if hasattr(self, 'assessment_policy_var') and self.assessment_policy_var.get():
                self._add_preview_section(content_container, "University Assessment Policies", 11, "bold")
                self._add_preview_text(content_container, "Requirements for make-up exams, assignments, and other work in this course are consistent with university policies that can be found in the Catalog.")

            # Add Instructions for Submitting Written Assignments section
            self._add_preview_section(content_container, "Instructions for Submitting Written Assignments", 12, "bold")
            self._add_preview_text(content_container, "All written assignments must be submitted as Word documents (.doc or .docx) through the \"Assignments\" portal in Canvas by the specified deadlines. Do NOT send assignments as PDF files.")
                        
            # Add Extensions & Make-Up Exams if checked
            if hasattr(self, 'extensions_policy_var') and self.extensions_policy_var.get():
                self._add_preview_section(content_container, "Extensions & Make-Up Exams", 11, "bold")
                self._add_preview_text(content_container, "Only the professor can authorize an extension or make-up exam, and all requests must be supported by documentation from a medical provider, Student Health Services, the Disability Resource Center, or the Dean of Students Office. Requirements for attendance and make-up exams, assignments, and other work in this course are consistent with university policies: https://catalog.ufl.edu/ugrad/current/regulations/info/attendance.aspx")
            
            # Add Late Submissions policy (show content from the input)
            self._add_preview_section(content_container, "Late Submissions", 11, "bold")
            if hasattr(self, 'late_policy_text') and self.late_policy_text.get("1.0", tk.END).strip():
                # Display the actual content from the late policy text input
                self._add_preview_text(content_container, self.late_policy_text.get("1.0", tk.END).strip())
            elif hasattr(self, 'late_policy_var') and self.late_policy_var.get() in self.late_policies:
                # If there's a selected policy from the dropdown but no custom text, use the dropdown's standard text
                selected_policy = self.late_policy_var.get()
                self._add_preview_text(content_container, self.late_policies[selected_policy])
            else:
                # Default text only if no input and no selection
                self._add_preview_text(content_container, "Late submission policy not specified.")
            
            # Add Extra Credit policy (using the exact same approach as Late Submissions)
            self._add_preview_section(content_container, "Extra Credit", 11, "bold")
            if hasattr(self, 'extra_credit_text') and self.extra_credit_text.get("1.0", tk.END).strip():
                # Display the actual content from the extra credit text input
                self._add_preview_text(content_container, self.extra_credit_text.get("1.0", tk.END).strip())
            elif hasattr(self, 'extra_credit_var') and self.extra_credit_var.get() in self.extra_credit_policies:
                # If there's a selected policy from the dropdown but no custom text, use the dropdown's standard text
                selected_policy = self.extra_credit_var.get()
                self._add_preview_text(content_container, self.extra_credit_policies[selected_policy])
            else:
                # Default text only if no input and no selection
                self._add_preview_text(content_container, "Extra credit policy not specified.")
            
            # Canvas Policy (always included)
            self._add_preview_section(content_container, "Canvas", 11, "bold")
            if hasattr(self, 'canvas_policy_text') and self.canvas_policy_text.get("1.0", tk.END).strip() != "-Replace with your Canvas Policies-":
                self._add_preview_text(content_container, self.canvas_policy_text.get("1.0", tk.END).strip())
            else:
                self._add_preview_text(content_container, canvas_policy_default)
            
            # Add Technology Policy
            self._add_preview_section(content_container, "Technology Policy", 11, "bold")
            self._add_preview_text(content_container, self.technology_policy_text.get("1.0", "end").strip())
            
            # Add Communication Policy
            self._add_preview_section(content_container, "Communication Policy", 11, "bold")
            self._add_preview_text(content_container, self.communication_policy_text.get("1.0", "end").strip())
            
            # Add Assignment Support section if enabled
            if self.outside_support_var.get():
                self._add_preview_section(content_container, "Assignment Support Outside the Classroom", 11, "bold")
                support_text = self.support_text.get("1.0", tk.END).strip()
                self._add_preview_text(content_container, support_text)
            
            # Add new "IV. Evaluations" section
            self._add_preview_section(content_container, "IV. Evaluations", 12, "bold")
            eval_text = (
                "UF course evaluation process\n"
                "Students are expected to provide professional and respectful feedback on the quality of "
                "instruction in this course by completing course evaluations online. Students can complete "
                "evaluations in three ways:\n\n"
                "1. The email they receive from GatorEvals,\n"
                "2. Their Canvas course menu under GatorEvals, or\n"
                "3. The central portal at https://my-ufl.bluera.com\n\n"
                "Guidance on how to provide constructive feedback is available at "
                "https://gatorevals.aa.ufl.edu/students/. Students will be notified when the evaluation "
                "period opens. Summaries of course evaluation results are available to students at "
                "https://gatorevals.aa.ufl.edu/public-results/."
            )
            self._add_preview_text(content_container, eval_text)

            # Add new "V. University Policies and Resources" section
            self._add_preview_section(content_container, "V. University Policies and Resources", 12, "bold")
            policies_text = (
                "Students requiring accommodation\n\n"
                "Students who experience learning barriers and would like to request academic "
                "accommodations should connect with the Disability Resource Center by visiting "
                "https://disability.ufl.edu/students/get-started/. It is important for students "
                "to share their accommodation letter with their instructor and discuss their "
                "access needs, as early as possible in the semester.\n\n"
                "University Honesty Policy\n\n"
                "UF students are bound by The Honor Pledge which states “We, the members of the "
                "University of Florida community, pledge to hold ourselves and our peers to the "
                "highest standards of honor and integrity by abiding by the Honor Code. On all "
                "work submitted for credit by students at the University of Florida, the "
                "following pledge is either required or implied: “On my honor, I have neither "
                "given nor received unauthorized aid in doing this assignment.” The Conduct Code "
                "specifies a number of behaviors that are in violation of this code and the "
                "possible sanctions. See the UF Conduct Code website for more information. If "
                "you have any questions or concerns, please consult with the instructor or TAs "
                "in this class."
            )
            self._add_preview_text(content_container, policies_text)
            
            # Move in-class recording policy here, after V. University Policies and Resources
            if hasattr(self, 'in_class_recording_var') and self.in_class_recording_var.get():
                self._add_preview_section(content_container, "In-class recording", 11, "bold")
                self._add_preview_text(content_container, recording_policy_default)
            if hasattr(self, 'conflict_resolution_var') and self.conflict_resolution_var.get():
                self._add_preview_section(content_container, "Procedure for conflict resolution", 11, "bold")
                self._add_preview_text(content_container, conflict_resolution_default)
            if self.campus_resources_var.get():
                self._add_preview_section(content_container, "Campus Resources", 11, "bold")
                self._add_preview_text(content_container, campus_resources_default)

            # Academic Resources
            if self.academic_resources_var.get():
                self._add_preview_section(content_container, "Academic Resources", 11, "bold")
                self._add_preview_text(content_container, academic_resources_default)
            
            # VI. Course Schedule
            self._add_preview_section(content_container, "VI. Calendar", 12, "bold")
            
            # Schedule Table Preview
            schedule_table_frame = ttk.Frame(content_container, style="Preview.TFrame")
            schedule_table_frame.pack(fill=tk.X, pady=5)
            
            # Header
            header_frame = ttk.Frame(schedule_table_frame, style="Preview.TFrame")
            header_frame.pack(fill=tk.X)
            ttk.Label(header_frame, text="Date", width=15, style="PreviewHeader.TLabel", borderwidth=1, relief="solid").pack(side=tk.LEFT)
            ttk.Label(header_frame, text="Topic", width=30, style="PreviewHeader.TLabel", borderwidth=1, relief="solid").pack(side=tk.LEFT)
            ttk.Label(header_frame, text="Readings/Preparation", width=50, style="PreviewHeader.TLabel", borderwidth=1, relief="solid").pack(side=tk.LEFT)
            ttk.Label(header_frame, text="Work Due", width=20, style="PreviewHeader.TLabel", borderwidth=1, relief="solid").pack(side=tk.LEFT)
            
            # Entries
            if hasattr(self, 'schedule_entries'):
                for entry in self.schedule_entries:
                    entry_frame = ttk.Frame(schedule_table_frame, style="Preview.TFrame")
                    entry_frame.pack(fill=tk.X)
                    ttk.Label(entry_frame, text=entry['date'].get(), width=15, style="PreviewCell.TLabel", borderwidth=1, relief="solid", wraplength=100).pack(side=tk.LEFT, fill=tk.Y)
                    ttk.Label(entry_frame, text=entry['topic'].get(), width=30, style="PreviewCell.TLabel", borderwidth=1, relief="solid", wraplength=200).pack(side=tk.LEFT, fill=tk.Y)
                    ttk.Label(entry_frame, text=entry['readings'].get("1.0", tk.END).strip(), width=50, style="PreviewCell.TLabel", borderwidth=1, relief="solid", wraplength=350).pack(side=tk.LEFT, fill=tk.Y)
                    ttk.Label(entry_frame, text=entry['work_due'].get(), width=20, style="PreviewCell.TLabel", borderwidth=1, relief="solid", wraplength=150).pack(side=tk.LEFT, fill=tk.Y)
            
            # Update scroll region after adding all content
            self.preview_content_frame.update_idletasks()
            self.preview_canvas.config(scrollregion=self.preview_canvas.bbox("all"))
        
        except Exception as e:
            print(f"Error updating document preview: {e}")
            import traceback
            traceback.print_exc()

        # At the very end of the function, after all content is added:
        try:
            # Add footer with page indicator
            footer_frame = ttk.Frame(content_container, style="Preview.TFrame")
            footer_frame.pack(fill=tk.X, pady=20)
            
            # Add note about pagination
            ttk.Label(footer_frame, 
                     text="Note: Page numbers will appear in the generated document",
                     font=("Times New Roman", 9, "italic"),
                     background="white").pack(side=tk.LEFT)
            
            # Show sample page number on right
            ttk.Label(footer_frame, 
                     text="Page X", 
                     font=("Times New Roman", 9),
                     background="white").pack(side=tk.RIGHT)
            
            # Add separator above footer
            ttk.Separator(content_container, orient=tk.HORIZONTAL).pack(
                fill=tk.X, pady=5, before=footer_frame)
        except Exception as e:
            print(f"Error adding page number preview: {e}")

    def _add_preview_section(self, parent, text, font_size=12, font_weight="normal"):
        """Add a section heading to the preview"""
        section = ttk.Label(parent, text=text,
                            font=("Times New Roman", font_size, font_weight),
                            background="white")
        section.pack(anchor="w", pady=5, fill=tk.X)
        return section

    def _add_preview_field(self, parent, label, value, indent=20):
        """Add a labeled field to the preview"""
        field_frame = ttk.Frame(parent, style="Preview.TFrame")
        field_frame.pack(fill=tk.X, padx=indent, pady=1)
        
        label_widget = ttk.Label(field_frame, text=label, 
                                 font=("Times New Roman", 10, "bold"),
                                 width=15, anchor="w",
                                 background="white")
        label_widget.pack(side=tk.LEFT)
        
        value_widget = ttk.Label(field_frame, text=value,
                                 font=("Times New Roman", 10),
                                 background="white",
                                 justify=tk.LEFT,
                                 wraplength=400)
        value_widget.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        return field_frame

    def _add_preview_text(self, parent, text, bold=False, indent=0, pady=2):
        """Add plain text to the preview"""
        text_widget = ttk.Label(parent, text=text,
                                font=("Times New Roman", 10, bold and "bold" or "normal"),
                                background="white",
                                justify=tk.LEFT,
                                wraplength=600)
        text_widget.pack(anchor="w", padx=indent, pady=pady, fill=tk.X)
        return text_widget

    def on_tab_changed(self, event):
        """Update preview when tab changes"""
        try:
            selected_tab = self.notebook.select()
            tab_id = self.notebook.index(selected_tab)
            # Document Preview tab
            if tab_id == 6:  # Document Preview tab
                self.update_document_preview()
            elif tab_id == 2:  # Learning Objectives tab
                self.update_lo_preview()
        except Exception as e:
            print(f"Error changing tabs: {e}")

    def add_outcome_entry(self, default_text=""):
        """Add a new Student Learning Outcome entry with a number"""
        # Create frame for the entry
        frame = ttk.Frame(self.outcome_entries_frame)
        frame.pack(fill=tk.X, pady=2)
        
        # Add number label for the entry
        number_label = ttk.Label(frame, text=f"{len(self.outcome_entries)+1}.", width=3)
        number_label.pack(side=tk.LEFT, padx=(5,0))
        
        # Add entry field
        entry = ttk.Entry(frame, width=60)
        entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        if default_text:
            entry.insert(0, default_text)
        
        def remove_outcome():
            frame.destroy()
            self.outcome_entries.remove(entry_dict)
            self.renumber_outcomes()
            self.update_outcomes_references()  # Update the references after removing
            self.update_lo_preview()
        
        remove_btn = ttk.Button(frame, text="X", 
                              command=remove_outcome,
                              style="Delete.TButton")
        remove_btn.pack(side=tk.LEFT, padx=5)
        
        # Create dictionary to store entry information
        entry_dict = {
            "frame": frame,
            "number": number_label,
            "entry": entry
        }
        
        # Add to list of entries
        self.outcome_entries.append(entry_dict)
        
        # Bind update event
        entry.bind("<KeyRelease>", lambda e: self.update_all_previews())
        
        # Update references immediately
        self.update_outcomes_references()
        
        return entry_dict

    def update_all_previews(self):
        """Update both previews"""
        self.update_outcomes_references()
        self.update_lo_preview()

    def renumber_outcomes(self):
        """Update the numbering of outcome entries after one is removed"""
        for i, entry in enumerate(self.outcome_entries):
            entry["number"].config(text=f"{i + 1}.")
        
        # Also update the preview
        self.update_lo_preview()

    def setup_styles(self):
        """Set up custom styles for the application"""
        style = ttk.Style()

        # General styles
        style.configure("TButton", padding=5, font=("Arial", 10))
        style.configure("TLabel", font=("Arial", 10))
        style.configure("TEntry", padding=5)
        style.configure("TFrame", background="#f5f5f5")
        
        # Custom styles
        style.configure("Heading.TLabel", font=("Arial", 12, "bold"))
        style.configure("Action.TButton", foreground="black", background="#0078D7", font=("Arial", 10, "bold"))
        style.map("Action.TButton", background=[("active", "#005A9E")])
        
        # Improve delete button appearance - changed foreground to black for better visibility
        style.configure("Delete.TButton", foreground="black", background="#D9534F", font=("Arial", 10, "bold"), 
                       padding=3, width=3)
        style.map("Delete.TButton", background=[("active", "#C9302C")])
        
        style.configure("Small.TButton", font=("Arial", 8))
        
        # Preview styles
        style.configure("Preview.TFrame", background="white")
        style.configure("Italic.TLabel", font=("Arial", 10, "italic"))

    def create_main_interface(self):
        """Create the main interface for the application"""
        # Create a main container frame
        self.main_container = ttk.Frame(self.root)
        self.main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Add a frame for the template dropdown at the top
        template_frame = ttk.Frame(self.main_container)
        template_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(template_frame, text="Select Template:", font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=(5, 10))
        self.template_var = tk.StringVar()
        template_options = ["Clear Template"] + self.template_names
        self.template_combo = ttk.Combobox(
            template_frame,
            textvariable=self.template_var,
            values=template_options,
            width=40
        )
        self.template_combo.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        self.template_combo.bind("<<ComboboxSelected>>", self.on_template_selected)

        # Create a notebook (tabbed interface)
        self.notebook = ttk.Notebook(self.main_container)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        # --- Add tabs in the order requested, with "Required" indicator ---
        # 1. Course Info (Required)
        self.create_course_info_tab()
        self.notebook.tab(0, text="★ Course Information [Required]")

        # 2. Instructor Info (Required)
        self.create_instructor_info_tab()
        self.notebook.tab(1, text="★ Instructor Information [Required]")

        # 3. Assignments & Grading
        self.create_assignments_tab()
        self.notebook.tab(2, text="★ Assignments & Grading [Required]")

        # 4. Schedule
        self.create_schedule_tab()
        self.notebook.tab(3, text="Calendar/Course Schedule")

        # 5. Learning Objectives
        self.create_learning_objectives_tab()
        self.notebook.tab(4, text="Learning Objectives & SLOs")

        # 6. Policies (all policy stuff at the end)
        self.create_policies_tab()
        self.notebook.tab(5, text="Policies")

        # 7. Document Preview
        self.create_document_preview_tab()
        self.notebook.tab(6, text="Document Preview")

        # Add action buttons at the bottom
        self.create_action_buttons()

    def add_objective_entry(self, default_text=""):
        """Add a new course objective entry with a number"""
        frame = ttk.Frame(self.objective_entries_frame)
        frame.pack(fill=tk.X, pady=2)
        
        number_label = ttk.Label(frame, text=f"{len(self.objective_entries)+1}.", width=3)
        number_label.pack(side=tk.LEFT, padx=(5,0))
        
        entry = ttk.Entry(frame, width=60)
        entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        if default_text:
            entry.insert(0, default_text)
        
        def remove_objective():
            frame.destroy()
            self.objective_entries.remove(obj_dict)
            self.renumber_objectives()
            self.update_document_preview()
        
        remove_btn = ttk.Button(frame, text="X", command=remove_objective, style="Delete.TButton")
        remove_btn.pack(side=tk.LEFT, padx=5)
        
        obj_dict = {"frame": frame, "number": number_label, "entry": entry}
        self.objective_entries.append(obj_dict)
        
        entry.bind("<KeyRelease>", lambda e: self.update_document_preview())
        
        return obj_dict

    def renumber_objectives(self):
        """Update numbering for course objectives"""
        for i, obj in enumerate(self.objective_entries):
            obj["number"].config(text=f"{i+1}.")
        
    def add_learning_objective_row(self, category="", slo="", assignments=""):
        """Add a new row to the Learning Objectives table"""
        row_frame = ttk.LabelFrame(self.lo_entries_frame, text=category or "New Category")
        row_frame.pack(fill=tk.X, pady=5, padx=5)
        
        # Default SLO text based on category
        default_slos = {
            "Content": "Identify, describe, and explain key themes, principles, and terminology; the history, theory and/or methodologies used; and social institutions, structures or processes.",
            "Critical Thinking": "Apply formal and informal qualitative or quantitative analysis effectively to examine the processes and means by which individuals make personal and group decisions. Assess and analyze ethical perspectives in individual and societal decisions.",
            "Communication": "Communication is the development and expression of ideas in written and oral forms."
        }
        
        # Category name
        name_frame = ttk.Frame(row_frame)
        name_frame.pack(fill=tk.X, pady=2)
        ttk.Label(name_frame, text="Category:").pack(side=tk.LEFT, padx=5)
        category_entry = ttk.Entry(name_frame, width=30)
        category_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        if category:
            category_entry.insert(0, category)
        
        # Social Science SLOs
        slo_frame = ttk.LabelFrame(row_frame, text="Social Science SLOs")
        slo_frame.pack(fill=tk.X, pady=2, padx=5)
        slo_text = scrolledtext.ScrolledText(slo_frame, width=40, height=4, wrap=tk.WORD)
        slo_text.pack(padx=5, pady=5)
        # Insert default SLO text based on category
        if category in default_slos:
            slo_text.insert("1.0", default_slos[category])
            slo_text.config(state="disabled")
        elif slo:
            slo_text.insert("1.0", slo)
        self.add_mousewheel_scrolling(slo_text)
        
        # State SLO Assignments
        assignments_frame = ttk.LabelFrame(row_frame, text="State SLO Assignments")
        assignments_frame.pack(fill=tk.X, pady=2, padx=5)
        assignments_text = scrolledtext.ScrolledText(assignments_frame, width=40, height=4, wrap=tk.WORD)
        assignments_text.pack(padx=5, pady=5)
        assignments_text.insert("1.0", assignments or self.get_outcomes_range())
        if category in default_slos:  # Only disable for predefined categories
            assignments_text.config(state="disabled")
        self.add_mousewheel_scrolling(assignments_text)
        
        # Course-Specific
        course_specific_frame = ttk.LabelFrame(row_frame, text="Course-Specific")
        course_specific_frame.pack(fill=tk.X, pady=2, padx=5)
        course_specific_text = scrolledtext.ScrolledText(course_specific_frame, width=40, height=4, wrap=tk.WORD)
        course_specific_text.pack(padx=5, pady=5)
        self.add_mousewheel_scrolling(course_specific_text)
        
        # Only show remove button for custom categories
        if category not in default_slos:
            ttk.Button(row_frame, text="Remove Category", 
                      command=lambda: self.remove_lo_category(row_frame, category),
                      style="Delete.TButton").pack(pady=5)
        
        # Store references
        self.learning_objectives_entries[category or f"category_{len(self.learning_objectives_entries)}"] = {
            "frame": row_frame,
            "category": category_entry,
            "slo": slo_text,
            "assignments": assignments_text,
            "course_specific": course_specific_text
        }
        
        # Bind updates only for editable fields
        for widget in [category_entry, course_specific_text]:
            widget.bind("<KeyRelease>", lambda e: self.update_lo_preview())
        
        return row_frame

    def update_lo_preview(self):
        """Update the Learning Objectives preview in the right panel"""
        if hasattr(self, 'preview_frame'):
            for widget in self.preview_frame.winfo_children():
                widget.destroy()
            
            # Title
            ttk.Label(self.preview_frame, 
                     text="Objectives—General Education and Social and Behavioral Sciences (S)", 
                     style="Heading.TLabel").pack(anchor="w", padx=5, pady=5)
            
            # Create table frame
            table_frame = ttk.Frame(self.preview_frame, relief=tk.SOLID, borderwidth=1)
            table_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            
            # Headers
            headers = ["CATEGORY", "SOCIAL SCIENCE SLOS", "STATE SLO ASSIGNMENTS", "COURSE-SPECIFIC"]
            for i, header in enumerate(headers):
                header_cell = ttk.Label(table_frame, text=header, 
                                      background="#808080", foreground="white",
                                      font=("Arial", 9, "bold"))
                header_cell.grid(row=0, column=i, sticky="nsew", padx=1, pady=1)
                table_frame.columnconfigure(i, weight=1 if i == 0 else 2)
            
            # Content rows
            row = 1
            for category, entries in self.learning_objectives_entries.items():
                if 'frame' in entries and not entries['frame'].winfo_exists():
                    continue
                
                # Get text from entries:
                category_text = entries['category'].get()
                slo_text = entries['slo'].get("1.0", tk.END).strip()
                assignments_text = entries['assignments'].get("1.0", tk.END).strip()
                course_specific_text = entries['course_specific'].get("1.0", tk.END).strip()
                
                # Create cells
                cells = [
                    (category_text, 120),
                    (slo_text, 200),
                    (assignments_text, 200),
                    (course_specific_text, 200)
                ]
                for col, (text, wraplength) in enumerate(cells):
                    cell = ttk.Label(table_frame, text=text, background="white",
                                   relief="solid", borderwidth=1, wraplength=wraplength,
                                   padding=5)
                    cell.grid(row=row, column=col, sticky="nsew", padx=1, pady=1)
                row += 1
            
            # Show placeholder if no entries
            if row == 1:
                ttk.Label(table_frame, text="No learning objectives defined",
                         style="Italic.TLabel").grid(row=1, column=0, columnspan=4, pady=10)

    def remove_lo_category(self, frame, category):
        """Remove a learning objective category"""
        if messagebox.askyesno("Confirm Delete", "Are you sure you want to remove this category?"):
            frame.destroy()
            if category in self.learning_objectives_entries:
                del self.learning_objectives_entries[category]
            self.update_lo_preview()

    def add_mousewheel_scrolling(self, widget, canvas=None):
        """Add mousewheel scrolling to a widget"""
        def _on_mousewheel(event):
            if canvas:
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            else:
                widget.yview_scroll(int(-1 * (event.delta / 120)), "units")

        # Bind mousewheel on Windows/MacOS
        widget.bind("<MouseWheel>", _on_mousewheel)
        # For Linux, bind Button-4 and Button-5
        widget.bind("<Button-4>", lambda e: _on_mousewheel(type('Event', (), {'delta': 120})))
        widget.bind("<Button-5>", lambda e: _on_mousewheel(type('Event', (), {'delta': -120})))

    def get_outcomes_range(self):
        """Get the range of outcomes based on current number of outcomes"""
        num_outcomes = len(self.outcome_entries)
        return f"Outcomes 1-{num_outcomes}" if num_outcomes > 0 else "No outcomes defined"

    def update_outcomes_references(self):
        """Update all references to outcomes in the learning objectives table"""
        outcomes_range = self.get_outcomes_range()
        for entries in self.learning_objectives_entries.values():
            if 'assignments' in entries and entries['assignments'].winfo_exists():
                current_state = entries['assignments'].cget('state')
                
                if current_state == 'disabled':
                    entries['assignments'].config(state='normal')
                current_text = entries['assignments'].get("1.0", tk.END).strip()
                if current_text.startswith("Outcome"):
                    entries['assignments'].delete("1.0", tk.END)
                    entries['assignments'].insert("1.0", outcomes_range)
                if current_state == 'disabled':
                    entries['assignments'].config(state='disabled')

    def parse_materials_markup(self, text, doc=None):
        """
        Parse Markdown-like markup in Required Materials and add to Word doc.
        - Supports the following syntax:
          - *italic* -> Text between single asterisks (*) will be italicized.
          - **bold** -> Text between double asterisks (**) will be bolded.
          - [text](url) -> Text inside square brackets ([text]) followed by a URL in parentheses (url) will become a hyperlink.
        """
        import re
        
        if doc is None:
            # Just return the text if no document is provided
            return text
        
        # Create paragraph for the content
        paragraph = doc.add_paragraph()
        
        # Regular expression to find markdown patterns
        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))')
        
        # Split the text based on the pattern
        parts = pattern.split(text)
        
        # Process each part
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('[') and '](' in part and part.endswith(')'):
                # Hyperlink
                # Extract the text and URL
                link_text = part[1:part.index('](')]
                url = part[part.index('](')+2:-1]
                
                # Add hyperlink
                run = paragraph.add_run(link_text)
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # Blue color
                run.font.underline = True
                
                # Add the actual hyperlink
                # This is a simplification - true hyperlinks need more work with XML
                # For now, we'll just make it look like a hyperlink
            else:
                # Regular text
                paragraph.add_run(part)
        
        return paragraph

    def show_formatting_help(self):
        """Display a popup with formatting help."""
        help_window = tk.Toplevel(self.root)
        help_window.title("Formatting Help")
        help_window.geometry("400x300")

        explanation = """
        Markdown-like Formatting Guide:
        - *italic* -> Text between single asterisks (*) will be italicized.
        - **bold** -> Text between double asterisks (**) will be bolded.
        - [text](url) -> Text inside square brackets ([text]) followed by a URL in parentheses (url) will become a hyperlink.

        Example:
        Input: "This is *italic*, **bold*, and [a link](http://example.com)."
        Output in Word:
          - "italic" will appear italicized.
          - "bold" will appear bolded.
          - "a link" will appear as a clickable hyperlink pointing to "http://example.com".
        """

        text_widget = tk.Text(help_window, wrap=tk.WORD, font=("Arial", 10))
        text_widget.insert(tk.END, explanation)
        text_widget.config(state=tk.DISABLED)  # Make the text read-only
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

if __name__ == "__main__":
    app = HistorySyllabusGenerator()
    app.run()
